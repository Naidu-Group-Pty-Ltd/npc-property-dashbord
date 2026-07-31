"""Reusable Word building blocks for the Aurixa Finance Portal templates.

Two layers:

* **Primitives** — thin wrappers over the raw WordprocessingML that
  ``python-docx`` does not expose (cell shading, per-edge borders, cell
  margins, character tracking, page fields).
* **Blocks** — the visual vocabulary shared by every template: cover panel,
  section band, guidance card, field grid, clause block, workflow ladder,
  signature panel. Every block reads its colour from a ``BrandProfile`` so a
  partner re-skin is a one-line change.

Everything is table-based. Word's table model is the only layout primitive
that survives the round trip through Word, LibreOffice, Google Docs and PDF
export with identical spacing, which is what "consistent across all documents"
has to mean in practice.
"""

from __future__ import annotations

from dataclasses import dataclass

from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Mm, Pt, RGBColor

from aurixa_brand import (
    BRAND_SLOTS,
    INSERT,
    LAYOUT,
    PALETTE,
    TYPE,
    BrandProfile,
    token,
)

CHECKBOX = "☐"  # ☐ BALLOT BOX
BULLET = "•"
ARROW = "›"


# ==========================================================================
# Primitives
# ==========================================================================

def _el(tag: str, **attrs) -> OxmlElement:
    node = OxmlElement(tag)
    for key, value in attrs.items():
        node.set(qn(f"w:{key}"), str(value))
    return node


#: WordprocessingML declares each property container as an ordered ``xsd:sequence``.
#: Word (and LibreOffice) reject a file whose children appear out of order, even
#: though the XML is well-formed, so every raw element we add has to be spliced
#: into the right slot rather than appended.
_CHILD_ORDER: dict[str, tuple[str, ...]] = {
    "rPr": (
        "rStyle", "rFonts", "b", "bCs", "i", "iCs", "caps", "smallCaps", "strike",
        "dstrike", "outline", "shadow", "emboss", "imprint", "noProof", "snapToGrid",
        "vanish", "webHidden", "color", "spacing", "w", "kern", "position", "sz",
        "szCs", "highlight", "u", "effect", "bdr", "shd", "fitText", "vertAlign",
        "rtl", "cs", "em", "lang", "eastAsianLayout", "specVanish", "oMath",
    ),
    "tcPr": (
        "cnfStyle", "tcW", "gridSpan", "hMerge", "vMerge", "tcBorders", "shd",
        "noWrap", "tcMar", "textDirection", "tcFitText", "vAlign", "hideMark",
    ),
    "tcBorders": ("top", "start", "left", "bottom", "end", "right", "insideH",
                  "insideV", "tl2br", "tr2bl"),
    "tblBorders": ("top", "start", "left", "bottom", "end", "right", "insideH",
                   "insideV"),
    "tblPr": (
        "tblStyle", "tblpPr", "tblOverlap", "bidiVisual", "tblStyleRowBandSize",
        "tblStyleColBandSize", "tblW", "jc", "tblCellSpacing", "tblInd",
        "tblBorders", "shd", "tblLayout", "tblCellMar", "tblLook", "tblCaption",
        "tblDescription",
    ),
    "trPr": (
        "cnfStyle", "divId", "gridBefore", "gridAfter", "wBefore", "wAfter",
        "cantSplit", "trHeight", "tblHeader", "tblCellSpacing", "jc", "hidden",
    ),
}


def _splice(parent, node) -> None:
    """Insert ``node`` into ``parent`` at its schema-mandated position."""
    container = parent.tag.split("}")[-1]
    order = _CHILD_ORDER.get(container)
    if order is None:
        parent.append(node)
        return
    name = node.tag.split("}")[-1]
    try:
        rank = order.index(name)
    except ValueError:
        parent.append(node)
        return
    for existing in parent:
        existing_name = existing.tag.split("}")[-1]
        if existing_name not in order or order.index(existing_name) > rank:
            existing.addprevious(node)
            return
    parent.append(node)


def _get_or_add(parent, tag: str):
    existing = parent.find(qn(tag))
    if existing is None:
        existing = OxmlElement(tag)
        _splice(parent, existing)
    return existing


def _replace(parent, tag: str, node) -> None:
    """Drop any existing ``tag`` children, then splice ``node`` into place."""
    for old in parent.findall(qn(tag)):
        parent.remove(old)
    _splice(parent, node)


def shade(cell, fill: str) -> None:
    """Apply a solid background fill to a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    _replace(tc_pr, "w:shd", _el("w:shd", val="clear", color="auto", fill=fill))


def cell_borders(cell, **edges) -> None:
    """Set per-edge borders.

    Each keyword is ``top`` / ``left`` / ``bottom`` / ``right`` (also
    ``insideH`` / ``insideV``) and takes either ``None`` for "no border" or a
    ``(size_eighths_pt, colour_hex)`` tuple. Omitted edges are left untouched.
    """
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = _get_or_add(tc_pr, "w:tcBorders")
    for edge, spec in edges.items():
        if spec is None:
            node = _el(f"w:{edge}", val="nil")
        else:
            size, colour = spec
            node = _el(f"w:{edge}", val="single", sz=size, space=0, color=colour)
        _replace(borders, f"w:{edge}", node)


def cell_margins(cell, top: int, left: int, bottom: int, right: int) -> None:
    """Set cell padding in dxa (twentieths of a point)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for tag, value in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        mar.append(_el(f"w:{tag}", w=value, type="dxa"))
    _replace(tc_pr, "w:tcMar", mar)


def valign(cell, value: str = "center") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    _replace(tc_pr, "w:vAlign", _el("w:vAlign", val=value))


def table_borders_none(table) -> None:
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        borders.append(_el(f"w:{edge}", val="nil"))
    _replace(table._tbl.tblPr, "w:tblBorders", borders)


def fixed_layout(table) -> None:
    _replace(table._tbl.tblPr, "w:tblLayout", _el("w:tblLayout", type="fixed"))
    table.autofit = False


def table_indent(table, dxa: int) -> None:
    _replace(table._tbl.tblPr, "w:tblInd", _el("w:tblInd", w=dxa, type="dxa"))


def row_height(row, mm: float, exact: bool = False) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    _replace(tr_pr, "w:trHeight",
             _el("w:trHeight", val=int(Mm(mm).twips), hRule="exact" if exact else "atLeast"))


def keep_row_together(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        _splice(tr_pr, OxmlElement("w:cantSplit"))


def repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        _splice(tr_pr, OxmlElement("w:tblHeader"))


def para_format(
    paragraph,
    *,
    before: float = 0,
    after: float = 0,
    line: float | None = None,
    align=None,
    keep_with_next: bool = False,
    keep_lines: bool = True,
    left_indent: float | None = None,
    hanging: float | None = None,
) -> None:
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line is not None:
        pf.line_spacing = line
    if align is not None:
        paragraph.alignment = align
    pf.keep_together = keep_lines
    pf.keep_with_next = keep_with_next
    if left_indent is not None:
        pf.left_indent = Pt(left_indent)
    if hanging is not None:
        pf.first_line_indent = Pt(-hanging)


def write(
    paragraph,
    text: str,
    *,
    font: str = TYPE.body,
    size: float = TYPE.body_text,
    bold: bool = False,
    italic: bool = False,
    colour: str = PALETTE.ink,
    caps: bool = False,
    tracking: float = 0.0,
    underline: bool = False,
):
    """Append a styled run. ``tracking`` is extra character spacing in points."""
    run = paragraph.add_run(text)
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.underline = underline
    run.font.color.rgb = RGBColor.from_string(colour)
    r_pr = run._r.get_or_add_rPr()
    # East-Asian + complex-script faces, so the font survives locale switches.
    fonts = _get_or_add(r_pr, "w:rFonts")
    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
        fonts.set(qn(f"w:{attr}"), font)
    if caps:
        _replace(r_pr, "w:caps", _el("w:caps", val="1"))
    if tracking:
        _replace(r_pr, "w:spacing", _el("w:spacing", val=int(round(tracking * 20))))
    return run


def para(container, text: str = "", **kwargs):
    """Add a paragraph to a document or cell and optionally write one run.

    Paragraph-level keywords (``before``, ``after``, ``line``, ``align``,
    ``keep_with_next``, ``left_indent``, ``hanging``) are separated from
    run-level keywords automatically.
    """
    para_keys = {"before", "after", "line", "align", "keep_with_next", "keep_lines",
                 "left_indent", "hanging"}
    p_kwargs = {k: v for k, v in kwargs.items() if k in para_keys}
    r_kwargs = {k: v for k, v in kwargs.items() if k not in para_keys}
    paragraph = container.add_paragraph()
    para_format(paragraph, **p_kwargs)
    if text:
        write(paragraph, text, **r_kwargs)
    return paragraph


def clear(cell):
    """Remove the empty paragraph python-docx seeds every new cell with."""
    for child in list(cell._tc):
        if child.tag == qn("w:p"):
            cell._tc.remove(child)
    return cell


def page_field(paragraph, instruction: str, **run_kwargs) -> None:
    """Insert a Word field such as ``PAGE`` or ``NUMPAGES``."""
    run = write(paragraph, "", **run_kwargs)
    begin = _el("w:fldChar", fldCharType="begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {instruction} "
    separate = _el("w:fldChar", fldCharType="separate")
    end = _el("w:fldChar", fldCharType="end")
    for node in (begin, instr, separate, end):
        run._r.append(node)


def page_break(container) -> None:
    paragraph = container.add_paragraph()
    para_format(paragraph, before=0, after=0)
    paragraph.add_run().add_break(WD_BREAK.PAGE)


def spacer(container, points: float = 6) -> None:
    paragraph = container.add_paragraph()
    para_format(paragraph, before=0, after=0)
    paragraph.paragraph_format.line_spacing = Pt(points)
    run = paragraph.add_run()
    run.font.size = Pt(1)


# ==========================================================================
# Table scaffolding
# ==========================================================================

def new_table(container, rows: int, cols: int, widths_mm: list[float]):
    """Create a fixed-layout, borderless table sized in millimetres.

    ``container`` may be the document, a cell, or a header/footer — the last two
    require an explicit total width, which ``Document.add_table`` rejects.
    """
    total = Mm(sum(widths_mm))
    try:
        table = container.add_table(rows=rows, cols=cols, width=total)
    except TypeError:
        table = container.add_table(rows=rows, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_borders_none(table)
    fixed_layout(table)
    table_indent(table, 0)
    for row in table.rows:
        keep_row_together(row)
        for index, cell in enumerate(row.cells):
            cell.width = Mm(widths_mm[index])
    for index, column in enumerate(table.columns):
        column.width = Mm(widths_mm[index])
    return table


def card(
    doc,
    *,
    fill: str,
    border: tuple[int, str] | None = None,
    accent_bar: tuple[int, str] | None = None,
    pad: tuple[int, int] | None = None,
    width_mm: float | None = None,
):
    """A single-cell panel. Returns the cell, ready for paragraphs.

    ``accent_bar`` draws a heavy left rule — the guidance-card signature.
    """
    width = width_mm or LAYOUT.content_width_mm
    table = new_table(doc, 1, 1, [width])
    cell = clear(table.cell(0, 0))
    shade(cell, fill)
    pad_y, pad_x = pad or (LAYOUT.card_pad_y, LAYOUT.card_pad_x)
    cell_margins(cell, pad_y, pad_x, pad_y, pad_x)
    edges: dict[str, tuple[int, str] | None] = {
        "top": border, "left": border, "bottom": border, "right": border,
    }
    if accent_bar:
        edges["left"] = accent_bar
    cell_borders(cell, **edges)
    return cell


# ==========================================================================
# Page furniture
# ==========================================================================

def configure_page(section) -> None:
    section.page_width = Mm(LAYOUT.page_width_mm)
    section.page_height = Mm(LAYOUT.page_height_mm)
    section.left_margin = Mm(LAYOUT.margin_side_mm)
    section.right_margin = Mm(LAYOUT.margin_side_mm)
    section.top_margin = Mm(LAYOUT.margin_top_mm)
    section.bottom_margin = Mm(LAYOUT.margin_bottom_mm)
    section.header_distance = Mm(LAYOUT.header_distance_mm)
    section.footer_distance = Mm(LAYOUT.footer_distance_mm)


def build_header(section, brand: BrandProfile, doc_title: str) -> None:
    """Running header: partner name left, document title right, gold hairline.

    Suppressed on page 1 so the cover panel reads full-bleed.
    """
    section.different_first_page_header_footer = True
    blank = section.first_page_header
    blank.is_linked_to_previous = False
    first_blank = blank.paragraphs[0]
    para_format(first_blank, before=0, after=0)
    first_blank.paragraph_format.line_spacing = Pt(1)
    first_blank.add_run().font.size = Pt(1)

    header = section.header
    header.is_linked_to_previous = False
    for stale in list(header.paragraphs)[1:]:
        stale._p.getparent().remove(stale._p)

    table = new_table(header, 1, 2, [LAYOUT.content_width_mm * 0.5, LAYOUT.content_width_mm * 0.5])
    left, right = (clear(c) for c in table.rows[0].cells)
    for cell in (left, right):
        cell_margins(cell, 0, 0, 70, 0)
        cell_borders(cell, bottom=(8, brand.accent), top=None, left=None, right=None)

    p = para(left, brand.company_name, size=TYPE.micro, bold=True, caps=True,
             tracking=1.1, colour=PALETTE.ink_soft, before=0, after=0)
    p = para(right, doc_title, size=TYPE.micro, bold=True, caps=True, tracking=1.1,
             colour=PALETTE.ink_soft, before=0, after=0, align=WD_ALIGN_PARAGRAPH.RIGHT)

    # Remove the seed paragraph Word keeps above the table.
    first = header.paragraphs[0]
    para_format(first, before=0, after=0)
    first.paragraph_format.line_spacing = Pt(1)
    first.add_run().font.size = Pt(1)


def build_footer(section, brand: BrandProfile) -> None:
    """Running footer: confidentiality + disclaimer left, version + page right."""
    section.different_first_page_header_footer = True
    for footer in (section.footer, section.first_page_footer):
        footer.is_linked_to_previous = False
        for stale in list(footer.paragraphs)[1:]:
            stale._p.getparent().remove(stale._p)

        table = new_table(footer, 1, 2,
                          [LAYOUT.content_width_mm * 0.52, LAYOUT.content_width_mm * 0.48])
        left, right = (clear(c) for c in table.rows[0].cells)
        for cell in (left, right):
            cell_margins(cell, 60, 0, 0, 0)
            cell_borders(cell, top=(4, PALETTE.line), bottom=None, left=None, right=None)

        p = para(left, brand.confidentiality, size=TYPE.micro, bold=True, caps=True,
                 tracking=1.0, colour=PALETTE.gold_dark, before=0, after=0)
        if brand.footer_disclaimer:
            write(p, f"   {brand.footer_disclaimer}", size=TYPE.micro,
                  colour=PALETTE.ink_faint)

        p = para(right, "", before=0, after=0, align=WD_ALIGN_PARAGRAPH.RIGHT)
        if brand.platform_note:
            write(p, f"{brand.platform_note}   |   ", size=TYPE.micro,
                  colour=PALETTE.ink_faint)
        write(p, f"v{brand.version}", size=TYPE.micro, bold=True, colour=PALETTE.ink_soft)
        write(p, "   |   Page ", size=TYPE.micro, colour=PALETTE.ink_faint)
        page_field(p, "PAGE", size=TYPE.micro, bold=True, colour=PALETTE.ink_soft)
        write(p, " of ", size=TYPE.micro, colour=PALETTE.ink_faint)
        page_field(p, "NUMPAGES", size=TYPE.micro, bold=True, colour=PALETTE.ink_soft)

        first = footer.paragraphs[0]
        para_format(first, before=0, after=0)
        first.paragraph_format.line_spacing = Pt(1)
        first.add_run().font.size = Pt(1)


# ==========================================================================
# Blocks
# ==========================================================================

def cover_panel(
    doc,
    brand: BrandProfile,
    *,
    title_lines: list[str],
    eyebrow: str,
    summary: str,
    chips: list[str],
    reference: str = "",
) -> None:
    """Full-bleed obsidian cover panel with logo slot, title and metadata.

    Followed by the legal caveat band and an issue-control strip, so the cover
    fills the page and carries everything a recipient needs to file the pack.
    """
    panel_table = new_table(doc, 1, 1, [LAYOUT.content_width_mm])
    row_height(panel_table.rows[0], 140.0)
    panel = clear(panel_table.cell(0, 0))
    shade(panel, brand.primary)
    cell_margins(panel, 430, 330, 430, 330)
    edge = (12, brand.accent)
    cell_borders(panel, top=edge, left=edge, bottom=edge, right=edge)
    valign(panel, "center")

    # Logo slot — dashed gold box, sized to a typical horizontal lockup.
    logo_table = new_table(panel, 1, 1, [62.0])
    logo_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    logo_cell = clear(logo_table.cell(0, 0))
    shade(logo_cell, brand.primary)
    cell_margins(logo_cell, 200, 120, 200, 120)
    dash = (8, brand.accent)
    cell_borders(logo_cell, top=dash, left=dash, bottom=dash, right=dash)
    para(logo_cell, brand.logo_placeholder, size=TYPE.micro, bold=True, caps=True,
         tracking=1.4, colour=brand.accent, align=WD_ALIGN_PARAGRAPH.CENTER,
         before=0, after=0)

    spacer(panel, 10)

    para(panel, eyebrow, font=TYPE.body, size=TYPE.cover_eyebrow, bold=True, caps=True,
         tracking=3.2, colour=brand.accent, align=WD_ALIGN_PARAGRAPH.CENTER,
         before=2, after=3)

    para(panel, brand.company_name, font=TYPE.body, size=TYPE.cover_eyebrow + 1.5,
         bold=True, caps=True, tracking=2.6, colour=PALETTE.paper,
         align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=8)

    for index, line in enumerate(title_lines):
        para(panel, line, font=TYPE.display, size=TYPE.cover_title, bold=True,
             tracking=0.4, colour=PALETTE.paper, align=WD_ALIGN_PARAGRAPH.CENTER,
             before=0, after=0 if index < len(title_lines) - 1 else 8,
             line=0.94, keep_with_next=True)

    # Gold rule under the title.
    rule = new_table(panel, 1, 1, [34.0])
    rule.alignment = WD_TABLE_ALIGNMENT.CENTER
    rule_cell = clear(rule.cell(0, 0))
    shade(rule_cell, brand.accent)
    cell_margins(rule_cell, 0, 0, 0, 0)
    p = para(rule_cell, "", before=0, after=0)
    p.paragraph_format.line_spacing = Pt(2)
    p.add_run().font.size = Pt(1)

    spacer(panel, 8)

    para(panel, summary, size=TYPE.cover_subtitle, colour="D8CEBE",
         align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=12, line=1.35)

    # Status chips.
    chip_width = LAYOUT.content_width_mm - 24
    chip_table = new_table(panel, 1, len(chips), [chip_width / len(chips)] * len(chips))
    for index, label in enumerate(chips):
        chip = clear(chip_table.cell(0, index))
        shade(chip, PALETTE.obsidian_soft)
        cell_margins(chip, 90, 60, 90, 60)
        cell_borders(chip, top=(4, brand.accent_deep), left=(4, brand.accent_deep),
                     bottom=(4, brand.accent_deep), right=(4, brand.accent_deep))
        para(chip, label, size=TYPE.micro, bold=True, caps=True, tracking=1.4,
             colour=brand.accent, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=0)

    spacer(panel, 10)

    meta = (f"VERSION {brand.version}   |   EFFECTIVE DATE: {brand.effective_date}"
            f"   |   {brand.confidentiality}")
    para(panel, meta, size=TYPE.cover_meta, bold=True, caps=True, tracking=1.6,
         colour="B9AC97", align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=0)

    # Caveat band directly beneath the panel.
    caveat = card(doc, fill=PALETTE.alert_soft, border=(4, "E8C4B4"),
                  pad=(110, 200))
    p = para(caveat, "IMPORTANT  ", size=TYPE.micro, bold=True, caps=True,
             tracking=1.2, colour=PALETTE.alert, before=0, after=0)
    write(p, brand.disclaimer, size=TYPE.body_small, colour=PALETTE.ink)

    spacer(doc, 7)

    # Issue-control strip — who prepared it, who it went to, and how to file it.
    field_grid(doc, brand, [
        Field("Prepared for", token("RECIPIENT ORGANISATION")),
        Field("Prepared by", token("SENDER NAME")),
        Field("Document reference", reference or token("REFERENCE")),
        Field("Date issued", token("DATE")),
    ])

    spacer(doc, 5)

    strip = card(doc, fill=brand.primary, border=(4, brand.primary), pad=(120, 200))
    p = para(strip, brand.company_name, size=TYPE.body_small, bold=True, caps=True,
             tracking=1.4, colour=brand.accent, before=0, after=2)
    para(strip, f"{brand.phone}   ·   {brand.email}   ·   {brand.website}   ·   {brand.address}",
         size=TYPE.micro, colour="C4B8A6", before=0, after=0)


def section_band(doc, brand: BrandProfile, number: str, title: str, kicker: str,
                 sub: str = "") -> None:
    """Numbered obsidian band that opens every section."""
    number_w = 18.0
    table = new_table(doc, 1, 2, [number_w, LAYOUT.content_width_mm - number_w])
    row = table.rows[0]
    row_height(row, 13.0)
    keep_row_together(row)

    badge = clear(row.cells[0])
    shade(badge, brand.accent)
    cell_margins(badge, 100, 60, 100, 60)
    valign(badge, "center")
    para(badge, number, font=TYPE.display, size=TYPE.section_number, bold=True,
         colour=brand.primary, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=0)

    body = clear(row.cells[1])
    shade(body, brand.primary)
    cell_margins(body, 100, 200, 100, 160)
    valign(body, "center")
    para(body, title, size=TYPE.section_title, bold=True, caps=True, tracking=1.8,
         colour=PALETTE.paper, before=0, after=1, keep_with_next=True)
    line = kicker if not sub else f"{kicker}   {ARROW}   {sub}"
    para(body, line, size=TYPE.section_kicker, caps=True, tracking=1.1,
         colour=brand.accent, before=0, after=0, keep_with_next=True)


def guidance_card(doc, brand: BrandProfile, label: str, text: str,
                  removable: bool = False) -> None:
    """Pale advisory card with a heavy gold left rule."""
    cell = card(doc, fill=brand.accent_pale, accent_bar=(24, brand.accent),
                pad=(130, 190))
    cell_borders(cell, top=(4, brand.accent_tint), bottom=(4, brand.accent_tint),
                 right=(4, brand.accent_tint))
    p = para(cell, label.upper(), size=TYPE.label, bold=True, caps=True, tracking=1.4,
             colour=PALETTE.gold_dark, before=0, after=2, keep_with_next=True)
    if removable:
        write(p, "   ·   DELETE BEFORE ISSUE", size=TYPE.micro, bold=True, caps=True,
              tracking=1.2, colour=PALETTE.alert)
    para(cell, text, size=TYPE.body_small, colour=PALETTE.ink, before=0, after=0,
         line=1.28)


def note_card(doc, brand: BrandProfile, label: str, text: str,
              tone: str = "neutral") -> None:
    """Principle / outcome card used to close a section."""
    tones = {
        "neutral": (PALETTE.sand, PALETTE.line, PALETTE.ink_soft),
        "brand": (PALETTE.violet_soft, "C9B8EA", PALETTE.violet),
        "info": (PALETTE.azure_soft, "BFE3F7", PALETTE.azure),
        "alert": (PALETTE.alert_soft, "E8C4B4", PALETTE.alert),
        "success": (PALETTE.success_soft, "B7E8CB", "15803D"),
    }
    fill, border, accent = tones.get(tone, tones["neutral"])
    cell = card(doc, fill=fill, border=(4, border), accent_bar=(18, accent),
                pad=(130, 190))
    para(cell, label.upper(), size=TYPE.label, bold=True, caps=True, tracking=1.4,
         colour=accent, before=0, after=2, keep_with_next=True)
    para(cell, text, size=TYPE.body_small, colour=PALETTE.ink, before=0, after=0,
         line=1.28)


@dataclass
class Field:
    """One label/value pair in a field grid."""

    label: str
    value: str = INSERT
    #: ``True`` renders the value as selectable checkbox options rather than a
    #: fill-in field, so the shading reads as "choose" not "type".
    choice: bool = False


def field_grid(doc, brand: BrandProfile, fields: list[Field], columns: int = 2) -> None:
    """Two-column (or one-column) label/value grid — the core form surface."""
    per_pair = LAYOUT.content_width_mm / columns
    label_w = per_pair * (0.42 if columns > 1 else 0.21)
    value_w = per_pair - label_w
    widths: list[float] = []
    for _ in range(columns):
        widths += [label_w, value_w]

    rows = (len(fields) + columns - 1) // columns
    table = new_table(doc, rows, columns * 2, widths)

    for index, entry in enumerate(fields):
        r, c = divmod(index, columns)
        row = table.rows[r]
        row_height(row, 8.6)
        keep_row_together(row)

        label_cell = clear(row.cells[c * 2])
        shade(label_cell, PALETTE.sand_deep)
        cell_margins(label_cell, LAYOUT.cell_pad_y, LAYOUT.cell_pad_x,
                     LAYOUT.cell_pad_y, 80)
        valign(label_cell, "center")
        cell_borders(label_cell, top=(4, PALETTE.paper_warm), left=None,
                     bottom=(4, PALETTE.paper_warm), right=None)
        para(label_cell, entry.label, size=TYPE.label, bold=True, caps=True,
             tracking=0.9, colour=PALETTE.ink, before=0, after=0, line=1.1)

        value_cell = clear(row.cells[c * 2 + 1])
        shade(value_cell, PALETTE.field if not entry.choice else PALETTE.paper_warm)
        cell_margins(value_cell, LAYOUT.cell_pad_y, LAYOUT.cell_pad_x,
                     LAYOUT.cell_pad_y, LAYOUT.cell_pad_x)
        valign(value_cell, "center")
        cell_borders(value_cell, top=(4, PALETTE.paper_warm), left=None,
                     bottom=(6, brand.accent_tint), right=None)
        colour = PALETTE.ink_faint if entry.value.startswith("<<") else PALETTE.ink
        para(value_cell, entry.value, size=TYPE.body_small, colour=colour,
             italic=entry.value.startswith("<<"), before=0, after=0, line=1.15)

    # Blank out any unused trailing cells so the grid stays rectangular.
    for index in range(len(fields), rows * columns):
        r, c = divmod(index, columns)
        for offset in (0, 1):
            blank = clear(table.rows[r].cells[c * 2 + offset])
            shade(blank, PALETTE.paper_warm)
            cell_margins(blank, LAYOUT.cell_pad_y, LAYOUT.cell_pad_x,
                         LAYOUT.cell_pad_y, LAYOUT.cell_pad_x)
            para(blank, "", before=0, after=0)


def responsibility_columns(doc, brand: BrandProfile, left: tuple[str, list[str]],
                           right: tuple[str, list[str]],
                           tones: tuple[str, str] = ("brand", "gold")) -> None:
    """Side-by-side responsibility / permission lists."""
    half = LAYOUT.content_width_mm / 2
    table = new_table(doc, 2, 2, [half, half])
    keep_row_together(table.rows[0])

    palettes = {
        "brand": (PALETTE.violet, PALETTE.violet_soft),
        "gold": (PALETTE.gold_dark, PALETTE.gold_tint),
        "alert": (PALETTE.alert, PALETTE.alert_soft),
        "success": ("15803D", PALETTE.success_soft),
    }

    for index, ((title, items), tone) in enumerate(zip((left, right), tones)):
        accent, tint = palettes.get(tone, palettes["gold"])

        head = clear(table.rows[0].cells[index])
        shade(head, accent)
        cell_margins(head, 110, 170, 110, 170)
        cell_borders(head, top=None, left=None, bottom=None,
                     right=(30, PALETTE.paper) if index == 0 else None)
        para(head, title, size=TYPE.label + 0.5, bold=True, caps=True, tracking=1.3,
             colour=PALETTE.paper, before=0, after=0, line=1.15, keep_with_next=True)

        body = clear(table.rows[1].cells[index])
        shade(body, tint)
        cell_margins(body, 140, 170, 140, 170)
        cell_borders(body, top=None, left=None, bottom=(4, PALETTE.line),
                     right=(30, PALETTE.paper) if index == 0 else None)
        for item_index, item in enumerate(items):
            p = para(body, "", before=0 if item_index == 0 else 3, after=0, line=1.25,
                     left_indent=9, hanging=9)
            write(p, f"{BULLET}  ", size=TYPE.body_small, bold=True, colour=accent)
            write(p, item, size=TYPE.body_small, colour=PALETTE.ink)


def clause_block(doc, brand: BrandProfile, heading: str, clauses: list[str]) -> None:
    """A numbered operative clause with hanging-indent sub-clauses."""
    cell = card(doc, fill=PALETTE.paper_warm, border=(4, PALETTE.line),
                accent_bar=(18, brand.primary), pad=(150, 190))
    para(cell, heading, size=TYPE.clause_heading, bold=True, colour=brand.primary,
         font=TYPE.display, before=0, after=5, keep_with_next=True)
    for index, text in enumerate(clauses):
        number, _, rest = text.partition(" ")
        p = para(cell, "", before=0 if index == 0 else 4, after=0, line=1.3,
                 left_indent=30, hanging=30)
        write(p, number, size=TYPE.body_text, bold=True, colour=PALETTE.gold_dark)
        write(p, f"\t{rest}", size=TYPE.body_text, colour=PALETTE.ink)
        p.paragraph_format.tab_stops.add_tab_stop(Pt(30))


def workflow_ladder(doc, brand: BrandProfile, steps: list[tuple[str, str, str]]) -> None:
    """Numbered stage ladder: ``(number, name, description)``."""
    num_w, name_w = 13.0, 38.0
    desc_w = LAYOUT.content_width_mm - num_w - name_w
    table = new_table(doc, len(steps), 3, [num_w, name_w, desc_w])

    for index, (number, name, description) in enumerate(steps):
        row = table.rows[index]
        row_height(row, 9.4)
        keep_row_together(row)
        tint = PALETTE.paper_warm if index % 2 == 0 else PALETTE.sand

        badge = clear(row.cells[0])
        shade(badge, brand.primary)
        cell_margins(badge, 90, 40, 90, 40)
        valign(badge, "center")
        cell_borders(badge, top=(4, PALETTE.paper_warm), bottom=(4, PALETTE.paper_warm),
                     left=None, right=None)
        para(badge, number, font=TYPE.display, size=TYPE.body_text + 1, bold=True,
             colour=brand.accent, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=0)

        name_cell = clear(row.cells[1])
        shade(name_cell, tint)
        cell_margins(name_cell, 90, 150, 90, 100)
        valign(name_cell, "center")
        cell_borders(name_cell, top=(4, PALETTE.paper_warm), bottom=(4, PALETTE.paper_warm),
                     left=None, right=None)
        para(name_cell, name, size=TYPE.label + 0.5, bold=True, caps=True, tracking=1.2,
             colour=PALETTE.gold_dark, before=0, after=0)

        desc_cell = clear(row.cells[2])
        shade(desc_cell, tint)
        cell_margins(desc_cell, 90, 100, 90, 150)
        valign(desc_cell, "center")
        cell_borders(desc_cell, top=(4, PALETTE.paper_warm), bottom=(4, PALETTE.paper_warm),
                     left=None, right=None)
        para(desc_cell, description, size=TYPE.body_small, colour=PALETTE.ink,
             before=0, after=0, line=1.2)


def data_table(doc, brand: BrandProfile, headers: list[str], rows: list[list[str]],
               widths_mm: list[float], zebra: bool = True) -> None:
    """Standard bordered data table with a branded header row."""
    table = new_table(doc, len(rows) + 1, len(headers), widths_mm)

    head_row = table.rows[0]
    row_height(head_row, 8.0)
    repeat_header(head_row)
    for index, label in enumerate(headers):
        cell = clear(head_row.cells[index])
        shade(cell, brand.primary)
        cell_margins(cell, 90, 120, 90, 120)
        valign(cell, "center")
        cell_borders(cell, top=None, left=None, bottom=(8, brand.accent), right=None)
        para(cell, label, size=TYPE.label, bold=True, caps=True, tracking=1.0,
             colour=PALETTE.paper, before=0, after=0, line=1.1)

    for r_index, values in enumerate(rows):
        row = table.rows[r_index + 1]
        row_height(row, 8.2)
        keep_row_together(row)
        tint = PALETTE.paper_warm if (r_index % 2 == 0 or not zebra) else PALETTE.sand
        for c_index, value in enumerate(values):
            cell = clear(row.cells[c_index])
            shade(cell, tint)
            cell_margins(cell, 85, 120, 85, 120)
            valign(cell, "center")
            cell_borders(cell, top=None, left=None, bottom=(4, PALETTE.line), right=None)
            placeholder = value.startswith("<<")
            para(cell, value, size=TYPE.body_small,
                 colour=PALETTE.ink_faint if placeholder else PALETTE.ink,
                 italic=placeholder, before=0, after=0, line=1.2)


def signature_panel(doc, brand: BrandProfile, blocks: list[tuple[str, list[str]]]) -> None:
    """Side-by-side execution blocks with ruled signature lines."""
    half = LAYOUT.content_width_mm / 2
    table = new_table(doc, 1, len(blocks), [LAYOUT.content_width_mm / len(blocks)] * len(blocks))
    keep_row_together(table.rows[0])

    for index, (title, lines) in enumerate(blocks):
        cell = clear(table.rows[0].cells[index])
        shade(cell, PALETTE.paper_warm)
        cell_margins(cell, 160, 190, 190, 190)
        cell_borders(cell, top=(12, brand.accent), left=(4, PALETTE.line),
                     bottom=(4, PALETTE.line),
                     right=(4, PALETTE.line))
        para(cell, title, size=TYPE.label + 0.5, bold=True, caps=True, tracking=1.3,
             colour=brand.primary, before=0, after=6, keep_with_next=True)
        for line_index, line in enumerate(lines):
            label, _, value = line.partition(":")
            p = para(cell, "", before=0 if line_index == 0 else 6, after=0, line=1.2)
            write(p, f"{label.strip()}   ", size=TYPE.micro, bold=True, caps=True,
                  tracking=1.0, colour=PALETTE.ink_soft)
            value = value.strip()
            if value:
                placeholder = value.startswith("<<")
                write(p, value, size=TYPE.body_small,
                      colour=PALETTE.ink_faint if placeholder else PALETTE.ink,
                      italic=placeholder)
            else:
                write(p, "_" * 26, size=TYPE.body_small, colour=PALETTE.line_strong)


def brand_slots_page(doc, brand: BrandProfile, extra_rows: list[tuple[str, str, str]] | None = None) -> None:
    """The white-label control sheet printed inside every template."""
    section_band(doc, brand, "W", "BRAND & CUSTOMISATION PANEL",
                 "White-label control sheet", "Complete once per organisation")
    spacer(doc, 6)
    guidance_card(
        doc, brand, "How to white-label this document",
        "Every item below appears in at least one place in this template. Replace the "
        "value once here, then use Word's Find & Replace (Ctrl+H / Cmd+H) on the merge "
        "token to update every occurrence in a single pass. Colour changes are applied "
        "from the Aurixa Finance Portal branding settings, or manually via the shading "
        "controls on each band. Delete this page before issuing the document.",
        removable=True,
    )
    spacer(doc, 6)
    data_table(
        doc, brand,
        ["Customisable area", "Merge token / location", "Notes"],
        [[name, location, note] for name, location, note in (extra_rows or []) + BRAND_SLOTS],
        [40.0, 52.0, LAYOUT.content_width_mm - 92.0],
    )
    spacer(doc, 6)
    note_card(
        doc, brand, "Colour system",
        "Primary (obsidian) drives section bands, clause rules and table headers. "
        "Accent (gold) drives numbers, chips, rules and field underlines. Semantic "
        "colours — green for confirmations, amber for caution, red for legal warnings — "
        "are fixed by design and should not be re-skinned, so a partner's palette never "
        "changes the meaning of a warning.",
        tone="brand",
    )


def document_map(doc, brand: BrandProfile, entries: list[tuple[str, str, str]]) -> None:
    """Contents page: ``(number, title, description)``."""
    section_band(doc, brand, "0", "DOCUMENT MAP", "What is inside this pack",
                 "Read before completing")
    spacer(doc, 6)
    data_table(
        doc, brand,
        ["#", "Section", "What it covers"],
        entries,
        [14.0, 54.0, LAYOUT.content_width_mm - 68.0],
    )


def set_core_properties(doc, brand: BrandProfile, title: str, subject: str,
                        keywords: str) -> None:
    props = doc.core_properties
    props.title = title
    props.subject = subject
    props.author = brand.company_name
    props.category = "Finance Portal template"
    props.comments = (
        "White-label template generated by scripts/finance-portal-templates. "
        "Edit the builder script and regenerate rather than editing this file by hand."
    )
    props.keywords = keywords


def base_document(brand: BrandProfile, doc_title: str):
    """A blank document with Aurixa page setup, header, footer and base styles."""
    from docx import Document

    doc = Document()
    configure_page(doc.sections[0])
    build_header(doc.sections[0], brand, doc_title)
    build_footer(doc.sections[0], brand)

    normal = doc.styles["Normal"]
    normal.font.name = TYPE.body
    normal.font.size = Pt(TYPE.body_text)
    normal.font.color.rgb = RGBColor.from_string(PALETTE.ink)
    r_pr = normal.element.get_or_add_rPr()
    fonts = _get_or_add(r_pr, "w:rFonts")
    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
        fonts.set(qn(f"w:{attr}"), TYPE.body)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.15
    return doc
