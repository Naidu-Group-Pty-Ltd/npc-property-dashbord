"""Aurixa Command Center component library.

Every template in the library is composed from these components and nothing
else. That is the constraint that makes 36 templates feel like one product:
a reader who has learned what a recommendation box looks like has learned it
everywhere, and a developer wiring content injection binds against a fixed
set of shapes rather than 36 bespoke layouts.

Flow, not canvas
----------------
Nothing here positions anything absolutely. Components are Word tables and
paragraphs in document flow, so injected content of any length grows the
component instead of overflowing it. The specific guarantees:

* Row heights are always ``atLeast``, never ``exact`` — a row that grows is
  correct, a row that clips is a defect.
* Every table row is ``cantSplit`` and every table header row is ``tblHeader``,
  so rows stay whole and headers repeat on every page a table spans.
* Section openers and block headings set ``keepNext``, so a heading is never
  stranded at the foot of a page.
* Tables are fixed-layout at explicit millimetre widths derived from the
  content width, so no table can grow past the margin.
* Vertical rhythm uses fixed-height spacer paragraphs rather than margin
  collapse, which Word, LibreOffice and Google Docs each resolve differently.

Signature convention
--------------------
Every component takes ``(container, theme, ...)``. ``container`` is a document
or a cell, so components nest. None of them return anything the caller must
position.
"""

from __future__ import annotations

import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "docgen"))
sys.path.insert(0, str(Path(__file__).parent))

from oxml import (  # noqa: E402
    Mm, Pt, WD_ALIGN_PARAGRAPH, WD_TABLE_ALIGNMENT,
    cell_borders, cell_margins, clear, horizontal_rule, keep_row_together,
    merge_across, new_table, page_break, page_field, para, para_format,
    repeat_header, row_height, set_style_font, shade, single_cell, spacer,
    tab_stop, table_caption, valign, write,
)
from theme import STATUS_TONES, Theme  # noqa: E402

BULLET = "•"
ARROW = "›"
CHECKBOX = "☐"
EM_DASH = "—"


# ==========================================================================
# Small helpers
# ==========================================================================

def _tone(theme: Theme, status: str) -> tuple[str, str, str, str]:
    """Resolve a status keyword to (fill, border, text, glyph) hex values."""
    fill_key, line_key, text_key, glyph = STATUS_TONES.get(
        status.strip().lower(), STATUS_TONES["n/a"])
    p = theme.palette
    return getattr(p, fill_key), getattr(p, line_key), getattr(p, text_key), glyph


def _is_token(value: str) -> bool:
    return isinstance(value, str) and value.strip().startswith("{{")


def _placeholder_colour(theme: Theme, value: str) -> str:
    return theme.palette.ink_faint if _is_token(value) else theme.palette.ink


def inner_width(theme: Theme, pad_dxa: int | None = None,
                allowance_mm: float = 1.4) -> float:
    """Usable width inside a padded card.

    Nesting a full-width table inside a padded cell overflows it, and Word
    silently centres the overflow instead of erroring — which reads as a
    misaligned indent. Every nested component sizes itself with this.
    """
    pad = pad_dxa if pad_dxa is not None else theme.card_pad[1]
    return theme.width - 2 * (pad / 20 * 25.4 / 72) - allowance_mm


def eyebrow(container, theme: Theme, text: str, *, colour: str | None = None,
            before: float = 0, after: float = 4, align=None):
    """The small tracked-out label that opens most components."""
    return para(container, text, font=theme.body, size=theme.type_scale.label,
                bold=True, caps=True, tracking=theme.family.label_tracking,
                colour=colour or theme.accent, before=before, after=after,
                align=align, keep_with_next=True)


def body(container, theme: Theme, text: str, *, size: float | None = None,
         colour: str | None = None, before: float = 0, after: float = 0,
         bold: bool = False, italic: bool = False, align=None):
    return para(container, text, font=theme.body,
                size=size or theme.type_scale.body,
                colour=colour or theme.palette.ink, bold=bold, italic=italic,
                before=before, after=after, line=theme.family.body_line,
                align=align)


def prose(container, theme: Theme, paragraphs: list[str], *, size: float | None = None,
          colour: str | None = None, first_gap: float = 0, gap: float = 5) -> None:
    """A run of body paragraphs. The workhorse for injected narrative content —
    executive summaries, rationale, recommendations, market commentary."""
    for index, text in enumerate(paragraphs):
        body(container, theme, text, size=size, colour=colour,
             before=first_gap if index == 0 else gap)


def bullets(container, theme: Theme, items: list[str], *, glyph: str = BULLET,
            colour: str | None = None, size: float | None = None,
            gap: float = 4, indent: float = 10) -> None:
    for index, item in enumerate(items):
        p = para(container, "", before=0 if index == 0 else gap, after=0,
                 line=theme.family.body_line, left_indent=indent, hanging=indent)
        write(p, f"{glyph}  ", font=theme.body, size=size or theme.type_scale.body,
              bold=True, colour=colour or theme.accent)
        write(p, item, font=theme.body, size=size or theme.type_scale.body,
              colour=theme.palette.ink)


def numbered(container, theme: Theme, items: list[str], *, start: int = 1,
             gap: float = 5, indent: float = 16) -> None:
    for index, item in enumerate(items):
        p = para(container, "", before=0 if index == 0 else gap, after=0,
                 line=theme.family.body_line, left_indent=indent, hanging=indent)
        write(p, f"{start + index}.", font=theme.body, size=theme.type_scale.body,
              bold=True, colour=theme.accent)
        write(p, f"\t{item}", font=theme.body, size=theme.type_scale.body,
              colour=theme.palette.ink)
        tab_stop(p, indent)


def rule(container, theme: Theme, *, colour: str | None = None, weight: int = 6,
         before: float = 0, after: float = 0) -> None:
    p = para(container, "", before=before, after=after)
    p.paragraph_format.line_spacing = Pt(1)
    p.add_run().font.size = Pt(1)
    horizontal_rule(p, colour or theme.palette.line, size=weight)


def gap(container, theme: Theme, multiplier: float = 1.0) -> None:
    spacer(container, theme.gap * multiplier)


# ==========================================================================
# Page furniture
# ==========================================================================

def configure_section(section, theme: Theme) -> None:
    g = theme.geometry
    section.page_width = Mm(g.page_width_mm)
    section.page_height = Mm(g.page_height_mm)
    section.left_margin = Mm(g.margin_side_mm)
    section.right_margin = Mm(g.margin_side_mm)
    section.top_margin = Mm(g.margin_top_mm)
    section.bottom_margin = Mm(g.margin_bottom_mm)
    section.header_distance = Mm(g.header_distance_mm)
    section.footer_distance = Mm(g.footer_distance_mm)


def _blank_first(part) -> None:
    for stale in list(part.paragraphs)[1:]:
        stale._p.getparent().remove(stale._p)
    first = part.paragraphs[0]
    para_format(first, before=0, after=0)
    first.paragraph_format.line_spacing = Pt(1)
    first.add_run().font.size = Pt(1)


def begin_landscape(doc, theme: Theme, doc_title: str) -> Theme:
    """Start a landscape section and return the theme sized for it.

    Word applies orientation per section, so a wide table needs its own section
    rather than a scaled-down font. The header and footer are rebuilt at the new
    width — inheriting the portrait ones would leave a 178mm rule on a 265mm
    page, which reads as a mistake rather than a design.
    """
    from docx.enum.section import WD_ORIENT, WD_SECTION

    landscape = theme.landscape
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.orientation = WD_ORIENT.LANDSCAPE
    configure_section(section, landscape)
    section.different_first_page_header_footer = False
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    _rebuild_running(section, landscape, doc_title)
    return landscape


def end_landscape(doc, theme: Theme, doc_title: str) -> None:
    """Return to portrait for the remainder of the document."""
    from docx.enum.section import WD_ORIENT, WD_SECTION

    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.orientation = WD_ORIENT.PORTRAIT
    configure_section(section, theme)
    section.different_first_page_header_footer = False
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    _rebuild_running(section, theme, doc_title)


def _rebuild_running(section, theme: Theme, doc_title: str) -> None:
    """Draw the running header and footer into a mid-document section."""
    header = section.header
    for stale in list(header.paragraphs)[1:]:
        stale._p.getparent().remove(stale._p)
    if theme.family.header_style != "none":
        half = theme.width / 2
        table = new_table(header, 1, 2, [half, half])
        left, right = (clear(c) for c in table.rows[0].cells)
        for cell in (left, right):
            cell_margins(cell, 0, 0, 70, 0)
            cell_borders(cell, bottom=(6, theme.accent), top=None, left=None, right=None)
        para(left, theme.brand.organisation_name, font=theme.body,
             size=theme.type_scale.micro, bold=True, caps=True,
             tracking=theme.family.label_tracking, colour=theme.palette.ink_soft,
             before=0, after=0)
        para(right, doc_title, font=theme.body, size=theme.type_scale.micro, bold=True,
             caps=True, tracking=theme.family.label_tracking,
             colour=theme.palette.ink_soft, before=0, after=0,
             align=WD_ALIGN_PARAGRAPH.RIGHT)
    _blank_first(header)

    footer = section.footer
    for stale in list(footer.paragraphs)[1:]:
        stale._p.getparent().remove(stale._p)
    table = new_table(footer, 1, 2, [theme.width * 0.56, theme.width * 0.44])
    left, right = (clear(c) for c in table.rows[0].cells)
    for cell in (left, right):
        cell_margins(cell, 70, 0, 40, 0)
        cell_borders(cell, top=(4, theme.palette.line), bottom=None, left=None, right=None)
    p = para(left, theme.brand.confidentiality, font=theme.body,
             size=theme.type_scale.micro, bold=True, caps=True,
             tracking=theme.family.label_tracking, colour=theme.accent, before=0, after=0)
    write(p, f"    {theme.brand.document_reference}", font=theme.body,
          size=theme.type_scale.micro, colour=theme.palette.ink_faint)
    p = para(right, "", before=0, after=0, align=WD_ALIGN_PARAGRAPH.RIGHT)
    if theme.brand.powered_by:
        write(p, f"{theme.brand.powered_by}    |    ", font=theme.body,
              size=theme.type_scale.micro, colour=theme.palette.ink_faint)
    write(p, f"v{theme.brand.version}", font=theme.body, size=theme.type_scale.micro,
          bold=True, colour=theme.palette.ink_soft)
    write(p, "    |    Page ", font=theme.body, size=theme.type_scale.micro,
          colour=theme.palette.ink_faint)
    page_field(p, "PAGE", font=theme.body, size=theme.type_scale.micro, bold=True,
               colour=theme.palette.ink_soft)
    write(p, " of ", font=theme.body, size=theme.type_scale.micro,
          colour=theme.palette.ink_faint)
    page_field(p, "NUMPAGES", font=theme.body, size=theme.type_scale.micro, bold=True,
               colour=theme.palette.ink_soft)
    _blank_first(footer)


def running_header(section, theme: Theme, doc_title: str) -> None:
    """Organisation left, document title right, family-appropriate rule.

    Suppressed on page 1 so the cover reads full-bleed.
    """
    section.different_first_page_header_footer = True
    blank = section.first_page_header
    blank.is_linked_to_previous = False
    _blank_first(blank)

    if theme.family.header_style == "none":
        return

    header = section.header
    header.is_linked_to_previous = False
    for stale in list(header.paragraphs)[1:]:
        stale._p.getparent().remove(stale._p)

    half = theme.width / 2
    table = new_table(header, 1, 2, [half, half])
    left, right = (clear(c) for c in table.rows[0].cells)
    banded = theme.family.header_style == "band"
    for cell in (left, right):
        cell_margins(cell, 60 if banded else 0, 90 if banded else 0, 70, 90 if banded else 0)
        if banded:
            shade(cell, theme.palette.cloud)
            cell_borders(cell, bottom=(6, theme.accent), top=None, left=None, right=None)
        elif theme.family.header_style == "rule":
            cell_borders(cell, bottom=(6, theme.accent), top=None, left=None, right=None)
        else:
            cell_borders(cell, bottom=(4, theme.palette.line), top=None, left=None, right=None)

    para(left, theme.brand.organisation_name, font=theme.body,
         size=theme.type_scale.micro, bold=True, caps=True,
         tracking=theme.family.label_tracking, colour=theme.palette.ink_soft,
         before=0, after=0)
    para(right, doc_title, font=theme.body, size=theme.type_scale.micro, bold=True,
         caps=True, tracking=theme.family.label_tracking,
         colour=theme.palette.ink_soft, before=0, after=0,
         align=WD_ALIGN_PARAGRAPH.RIGHT)
    _blank_first(header)


def running_footer(section, theme: Theme) -> None:
    """Confidentiality and reference left; attribution, version and page right.

    Written into both the first-page and default footers, so the cover still
    carries document control.
    """
    section.different_first_page_header_footer = True
    banded = theme.family.footer_style == "band"

    for footer in (section.footer, section.first_page_footer):
        footer.is_linked_to_previous = False
        for stale in list(footer.paragraphs)[1:]:
            stale._p.getparent().remove(stale._p)

        table = new_table(footer, 1, 2, [theme.width * 0.56, theme.width * 0.44])
        left, right = (clear(c) for c in table.rows[0].cells)
        for cell in (left, right):
            cell_margins(cell, 70, 90 if banded else 0, 40, 90 if banded else 0)
            if banded:
                shade(cell, theme.palette.mist)
                cell_borders(cell, top=(6, theme.accent), bottom=None, left=None, right=None)
            elif theme.family.footer_style == "rule":
                cell_borders(cell, top=(4, theme.palette.line), bottom=None,
                             left=None, right=None)
            else:
                cell_borders(cell, top=None, bottom=None, left=None, right=None)

        p = para(left, theme.brand.confidentiality, font=theme.body,
                 size=theme.type_scale.micro, bold=True, caps=True,
                 tracking=theme.family.label_tracking, colour=theme.accent,
                 before=0, after=0)
        write(p, f"    {theme.brand.document_reference}", font=theme.body,
              size=theme.type_scale.micro, colour=theme.palette.ink_faint)

        p = para(right, "", before=0, after=0, align=WD_ALIGN_PARAGRAPH.RIGHT)
        if theme.brand.powered_by:
            write(p, f"{theme.brand.powered_by}    |    ", font=theme.body,
                  size=theme.type_scale.micro, colour=theme.palette.ink_faint)
        write(p, f"v{theme.brand.version}", font=theme.body,
              size=theme.type_scale.micro, bold=True, colour=theme.palette.ink_soft)
        write(p, "    |    Page ", font=theme.body, size=theme.type_scale.micro,
              colour=theme.palette.ink_faint)
        page_field(p, "PAGE", font=theme.body, size=theme.type_scale.micro,
                   bold=True, colour=theme.palette.ink_soft)
        write(p, " of ", font=theme.body, size=theme.type_scale.micro,
              colour=theme.palette.ink_faint)
        page_field(p, "NUMPAGES", font=theme.body, size=theme.type_scale.micro,
                   bold=True, colour=theme.palette.ink_soft)
        _blank_first(footer)


# ==========================================================================
# Covers
# ==========================================================================

def _logo_slot(container, theme: Theme, text: str, width_mm: float, *,
               on_dark: bool, dashed: bool = True) -> None:
    table = new_table(container, 1, 1, [width_mm])
    cell = clear(table.cell(0, 0))
    frame = theme.accent if on_dark else theme.palette.line_strong
    style = "dashed" if dashed else "single"
    cell_margins(cell, 190, 130, 190, 130)
    cell_borders(cell, top=(8, frame, style), left=(8, frame, style),
                 bottom=(8, frame, style), right=(8, frame, style))
    para(cell, text, font=theme.body, size=theme.type_scale.micro, bold=True,
         caps=True, tracking=theme.family.label_tracking + 0.4,
         colour=theme.accent if on_dark else theme.palette.ink_soft,
         align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=0)


def cover(container, theme: Theme, *, title: str, subtitle: str = "",
          eyebrow_text: str = "", chips: list[str] | None = None,
          prepared_for: bool = True, image_caption: str = "") -> None:
    """Cover page. Architecture varies by design family; the information
    hierarchy does not — eyebrow, organisation, title, subtitle, then document
    control. Every variant ends with the same issue-control block so a
    recipient can file any document in the library the same way."""
    style = theme.family.cover_style
    p = theme.palette
    ts = theme.type_scale

    if style in ("band", "panel", "fullbleed"):
        _cover_panel(container, theme, title, subtitle, eyebrow_text, chips,
                     image_caption)
    elif style == "split":
        _cover_split(container, theme, title, subtitle, eyebrow_text, chips)
    elif style == "editorial":
        _cover_editorial(container, theme, title, subtitle, eyebrow_text,
                         image_caption)
    else:
        _cover_minimal(container, theme, title, subtitle, eyebrow_text)

    if prepared_for:
        spacer(container, 8)
        document_control(container, theme, compact=True)


def _cover_panel(container, theme, title, subtitle, eyebrow_text, chips,
                 image_caption) -> None:
    p, ts = theme.palette, theme.type_scale
    fullbleed = theme.family.cover_style == "fullbleed"

    if fullbleed and theme.family.cover_image_slot:
        frame = single_cell(container, theme.width, fill=p.cloud,
                            border=(6, p.line_strong), pad=(1150, 200))
        para(frame, theme.brand.cover_image_placeholder, font=theme.body,
             size=ts.label, bold=True, caps=True, tracking=1.6,
             colour=p.ink_faint, align=WD_ALIGN_PARAGRAPH.CENTER,
             before=0, after=0)
        if image_caption:
            para(frame, image_caption, font=theme.body, size=ts.micro,
                 colour=p.ink_faint, align=WD_ALIGN_PARAGRAPH.CENTER,
                 before=4, after=0)
        spacer(container, 0)

    table = new_table(container, 1, 1, [theme.width])
    row_height(table.rows[0], 118.0 if fullbleed else 138.0)
    panel = clear(table.cell(0, 0))
    shade(panel, theme.primary)
    cell_margins(panel, 380, 320, 380, 320)
    cell_borders(panel, top=(10, theme.accent), left=None, bottom=None, right=None)
    valign(panel, "center")

    _logo_slot(panel, theme, theme.brand.logo_placeholder, 62.0, on_dark=True)
    spacer(panel, 10)

    if eyebrow_text:
        para(panel, eyebrow_text, font=theme.body, size=ts.cover_eyebrow, bold=True,
             caps=True, tracking=theme.family.label_tracking + 1.6,
             colour=theme.accent, before=0, after=4)
    para(panel, theme.brand.organisation_name, font=theme.body,
         size=ts.cover_eyebrow + 1.5, bold=True, caps=True, tracking=2.4,
         colour=p.ink_invert, before=0, after=9)
    para(panel, title, font=theme.display,
         size=ts.cover_title if len(title) < 44 else ts.cover_title_sm,
         bold=True, tracking=theme.family.display_tracking,
         colour=p.ink_invert, before=0, after=8, line=0.96, keep_with_next=True)
    if theme.family.cover_rule:
        bar = new_table(panel, 1, 1, [30.0])
        bar.alignment = WD_TABLE_ALIGNMENT.LEFT
        bar_cell = clear(bar.cell(0, 0))
        shade(bar_cell, theme.accent)
        cell_margins(bar_cell, 0, 0, 0, 0)
        q = para(bar_cell, "", before=0, after=0)
        q.paragraph_format.line_spacing = Pt(2.4)
        q.add_run().font.size = Pt(1)
        spacer(panel, 8)
    if subtitle:
        para(panel, subtitle, font=theme.body, size=ts.cover_subtitle,
             colour=p.ink_invert_soft, before=0, after=0, line=1.34)
    if chips:
        spacer(panel, 9)
        _chip_row(panel, theme, chips, on_dark=True)


def _cover_split(container, theme, title, subtitle, eyebrow_text, chips) -> None:
    p, ts = theme.palette, theme.type_scale
    left_w = theme.width * 0.40
    table = new_table(container, 1, 2, [left_w, theme.width - left_w])
    row_height(table.rows[0], 150.0)

    left = clear(table.rows[0].cells[0])
    shade(left, theme.primary)
    cell_margins(left, 340, 240, 340, 240)
    valign(left, "top")
    _logo_slot(left, theme, theme.brand.logo_placeholder, left_w - 24, on_dark=True)
    spacer(left, 14)
    para(left, theme.brand.organisation_name, font=theme.body, size=ts.label,
         bold=True, caps=True, tracking=2.0, colour=theme.accent, before=0, after=5)
    para(left, theme.brand.tagline, font=theme.body, size=ts.micro,
         colour=p.ink_invert_soft, before=0, after=0, line=1.4)
    spacer(left, 12)
    for label_text, value in (("Prepared for", theme.brand.client_name),
                              ("Prepared by", theme.brand.author_name),
                              ("Issued", theme.brand.issue_date)):
        para(left, label_text.upper(), font=theme.body, size=ts.micro, bold=True,
             caps=True, tracking=1.4, colour=theme.accent, before=6, after=1)
        para(left, value, font=theme.body, size=ts.body_sm,
             colour=p.ink_invert, before=0, after=0)

    right = clear(table.rows[0].cells[1])
    shade(right, p.paper)
    cell_margins(right, 340, 260, 340, 100)
    valign(right, "center")
    if eyebrow_text:
        para(right, eyebrow_text, font=theme.body, size=ts.cover_eyebrow, bold=True,
             caps=True, tracking=theme.family.label_tracking + 1.2,
             colour=theme.accent, before=0, after=6)
    para(right, title, font=theme.display, size=ts.cover_title,
         bold=True, tracking=theme.family.display_tracking, colour=theme.primary,
         before=0, after=8, line=1.0, keep_with_next=True)
    rule(right, theme, colour=theme.accent, weight=12, after=8)
    if subtitle:
        para(right, subtitle, font=theme.body, size=ts.cover_subtitle,
             colour=p.ink_soft, before=0, after=0, line=1.4)
    if chips:
        spacer(right, 10)
        _chip_row(right, theme, chips, on_dark=False, width=theme.width - left_w - 18)


def _cover_editorial(container, theme, title, subtitle, eyebrow_text,
                     image_caption) -> None:
    p, ts = theme.palette, theme.type_scale
    frame = single_cell(container, theme.width, fill=p.cloud,
                        border=(6, p.line_strong), pad=(1250, 200))
    para(frame, theme.brand.cover_image_placeholder, font=theme.body, size=ts.label,
         bold=True, caps=True, tracking=1.8, colour=p.ink_faint,
         align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=0)
    if image_caption:
        para(frame, image_caption, font=theme.body, size=ts.micro, italic=True,
             colour=p.ink_faint, align=WD_ALIGN_PARAGRAPH.CENTER, before=5, after=0)

    spacer(container, 16)
    _logo_slot(container, theme, theme.brand.logo_placeholder, 68.0, on_dark=False)
    spacer(container, 14)
    if eyebrow_text:
        para(container, eyebrow_text, font=theme.body, size=ts.cover_eyebrow,
             bold=True, caps=True, tracking=theme.family.label_tracking + 2.0,
             colour=theme.accent, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=8)
    para(container, title, font=theme.display, size=ts.cover_title + 2,
         bold=False, tracking=theme.family.display_tracking, colour=theme.primary,
         align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=10, line=1.02,
         keep_with_next=True)
    rule(container, theme, colour=theme.accent, weight=8, after=10)
    if subtitle:
        para(container, subtitle, font=theme.body, size=ts.cover_subtitle,
             colour=p.ink_soft, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=0,
             line=1.5)


def _cover_minimal(container, theme, title, subtitle, eyebrow_text) -> None:
    p, ts = theme.palette, theme.type_scale
    spacer(container, 6)
    header = new_table(container, 1, 2, [theme.width * 0.5, theme.width * 0.5])
    left, right = (clear(c) for c in header.rows[0].cells)
    for cell in (left, right):
        cell_margins(cell, 0, 0, 140, 0)
        cell_borders(cell, bottom=(12, theme.primary), top=None, left=None, right=None)
    para(left, theme.brand.organisation_name, font=theme.body, size=ts.label,
         bold=True, caps=True, tracking=1.6, colour=theme.primary, before=0, after=0)
    para(right, theme.brand.logo_placeholder, font=theme.body, size=ts.micro,
         caps=True, tracking=1.4, colour=p.ink_faint, before=0, after=0,
         align=WD_ALIGN_PARAGRAPH.RIGHT)

    spacer(container, 18)
    if eyebrow_text:
        para(container, eyebrow_text, font=theme.body, size=ts.cover_eyebrow,
             bold=True, caps=True, tracking=theme.family.label_tracking + 1.0,
             colour=theme.accent, before=0, after=6)
    para(container, title, font=theme.display, size=ts.cover_title_sm, bold=True,
         colour=theme.primary, before=0, after=6, line=1.06, keep_with_next=True)
    if subtitle:
        para(container, subtitle, font=theme.body, size=ts.cover_subtitle,
             colour=p.ink_soft, before=0, after=0, line=1.36)


def _chip_row(container, theme, chips: list[str], *, on_dark: bool,
              width: float | None = None) -> None:
    total = width or theme.width - 26
    table = new_table(container, 1, len(chips), [total / len(chips)] * len(chips))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT if not on_dark else WD_TABLE_ALIGNMENT.CENTER
    for index, text in enumerate(chips):
        chip = clear(table.cell(0, index))
        shade(chip, theme.palette.navy_mid if on_dark else theme.accent_tint)
        cell_margins(chip, 90, 70, 90, 70)
        edge = (4, theme.accent if on_dark else theme.palette.line_strong)
        cell_borders(chip, top=edge, left=edge, bottom=edge, right=edge)
        para(chip, text, font=theme.body, size=theme.type_scale.micro, bold=True,
             caps=True, tracking=1.4,
             colour=theme.accent if on_dark else theme.primary,
             align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=0)


# ==========================================================================
# Section structure
# ==========================================================================

def section_opener(container, theme: Theme, number: str, title: str,
                   kicker: str = "") -> None:
    """Opens every numbered section. Four family variants, same hierarchy."""
    style = theme.family.section_style
    p, ts = theme.palette, theme.type_scale

    if style == "bar":
        chip_w = 17.0 if theme.family.section_number_chip else 0.0
        cols = [chip_w, theme.width - chip_w] if chip_w else [theme.width]
        table = new_table(container, 1, len(cols), cols)
        row_height(table.rows[0], 12.0)
        cells = list(table.rows[0].cells)
        if chip_w:
            chip = clear(cells[0])
            shade(chip, theme.accent)
            cell_margins(chip, 90, 50, 90, 50)
            valign(chip, "center")
            para(chip, number, font=theme.display, size=ts.h2 + 1, bold=True,
                 colour=p.ink_invert, align=WD_ALIGN_PARAGRAPH.CENTER,
                 before=0, after=0, keep_with_next=True)
        bar = clear(cells[-1])
        shade(bar, theme.primary)
        cell_margins(bar, 95, 190, 95, 150)
        valign(bar, "center")
        para(bar, title, font=theme.display, size=ts.h2, bold=True, caps=True,
             tracking=theme.family.label_tracking + 0.4, colour=p.ink_invert,
             before=0, after=1 if kicker else 0, keep_with_next=True)
        if kicker:
            para(bar, kicker, font=theme.body, size=ts.section_kicker if hasattr(
                ts, "section_kicker") else ts.micro, caps=True, tracking=1.1,
                colour=theme.accent, before=0, after=0, keep_with_next=True)

    elif style == "tab":
        table = new_table(container, 2, 1, [theme.width])
        tab = clear(table.rows[0].cells[0])
        shade(tab, theme.accent)
        cell_margins(tab, 60, 150, 60, 150)
        para(tab, f"{number}   {title}".strip(), font=theme.body, size=ts.label + 0.5,
             bold=True, caps=True, tracking=theme.family.label_tracking + 0.4,
             colour=p.ink_invert, before=0, after=0, keep_with_next=True)
        base = clear(table.rows[1].cells[0])
        shade(base, theme.palette.mist)
        cell_margins(base, 110, 150, 110, 150)
        cell_borders(base, bottom=(4, p.line), left=None, right=None, top=None)
        para(base, kicker or title, font=theme.display, size=ts.h2, bold=True,
             colour=theme.primary, before=0, after=0, keep_with_next=True)

    elif style == "numbered":
        p_head = para(container, "", before=0, after=3, keep_with_next=True)
        if number:
            write(p_head, f"{number}   ", font=theme.display, size=ts.h1,
                  bold=False, colour=theme.accent,
                  tracking=theme.family.display_tracking)
        write(p_head, title, font=theme.display, size=ts.h1, bold=False,
              colour=theme.primary, tracking=theme.family.display_tracking)
        if kicker:
            para(container, kicker, font=theme.body, size=ts.label, bold=True,
                 caps=True, tracking=theme.family.label_tracking,
                 colour=p.ink_soft, before=0, after=5, keep_with_next=True)
        rule(container, theme, colour=theme.accent, weight=10, after=0)

    else:  # "rule" / "plain"
        p_head = para(container, "", before=0, after=0, keep_with_next=True)
        if number and theme.family.section_number_chip:
            write(p_head, f"{number}  ", font=theme.body, size=ts.h2, bold=True,
                  colour=theme.accent)
        write(p_head, title, font=theme.display, size=ts.h2, bold=True, caps=True,
              tracking=theme.family.label_tracking * 0.6, colour=theme.primary)
        horizontal_rule(p_head, theme.accent, size=10, space=4)
        if kicker:
            para(container, kicker, font=theme.body, size=ts.caption,
                 colour=p.ink_soft, before=5, after=0, keep_with_next=True)


def subsection(container, theme: Theme, title: str, *, before: float = 8) -> None:
    para(container, title, font=theme.display, size=theme.type_scale.h3, bold=True,
         colour=theme.primary, before=before, after=4, keep_with_next=True)


def run_in_heading(container, theme: Theme, title: str, text: str,
                   *, before: float = 6) -> None:
    p = para(container, "", before=before, after=0, line=theme.family.body_line)
    write(p, f"{title}   ", font=theme.body, size=theme.type_scale.h4, bold=True,
          colour=theme.primary)
    write(p, text, font=theme.body, size=theme.type_scale.body,
          colour=theme.palette.ink)


def divider(container, theme: Theme, label_text: str = "") -> None:
    """Full-width section divider. Used between major parts of long reports."""
    cell = single_cell(container, theme.width, fill=theme.palette.mist,
                       border=(4, theme.palette.line),
                       accent_bar=(20, theme.accent), pad=theme.tight_pad)
    para(cell, label_text or EM_DASH, font=theme.body, size=theme.type_scale.label,
         bold=True, caps=True, tracking=theme.family.label_tracking + 0.6,
         colour=theme.palette.ink_soft, before=0, after=0)


# ==========================================================================
# Panels, cards and callouts
# ==========================================================================

def _card_frame(container, theme: Theme, *, tone_fill: str, tone_line: str,
                accent: str, width: float | None = None):
    style = theme.family.card_style
    width = width or theme.width
    if style == "plain":
        return single_cell(container, width, fill=None, border=None,
                           accent_bar=(14, accent), pad=theme.card_pad)
    if style == "outlined":
        return single_cell(container, width, fill=theme.palette.paper,
                           border=(6, tone_line), accent_bar=(18, accent),
                           pad=theme.card_pad)
    if style == "shadowline":
        cell = single_cell(container, width, fill=tone_fill, border=(4, tone_line),
                           pad=theme.card_pad)
        cell_borders(cell, bottom=(14, tone_line))
        return cell
    return single_cell(container, width, fill=tone_fill, border=(4, tone_line),
                       accent_bar=(18, accent), pad=theme.card_pad)


def executive_summary(container, theme: Theme, *, paragraphs: list[str],
                      headline: str = "", takeaways: list[str] | None = None,
                      title: str = "Executive summary") -> None:
    """The most-read component in the library. Headline claim, supporting
    narrative of any length, then the three-to-five things a reader must
    remember. Grows without limit; the takeaway list is optional."""
    cell = _card_frame(container, theme, tone_fill=theme.palette.mist,
                       tone_line=theme.palette.line, accent=theme.accent)
    eyebrow(cell, theme, title)
    if headline:
        para(cell, headline, font=theme.display, size=theme.type_scale.h3 + 1,
             bold=True, colour=theme.primary, before=0, after=6,
             line=1.22, keep_with_next=True)
    prose(cell, theme, paragraphs, gap=5)
    if takeaways:
        spacer(cell, 8)
        eyebrow(cell, theme, "What this means", colour=theme.palette.ink_soft, after=4)
        bullets(cell, theme, takeaways, glyph=ARROW)


def highlight_box(container, theme: Theme, *, title: str, text: str,
                  tone: str = "info", items: list[str] | None = None) -> None:
    fill, line, ink, _ = _tone(theme, tone)
    cell = _card_frame(container, theme, tone_fill=fill, tone_line=line, accent=ink)
    eyebrow(cell, theme, title, colour=ink)
    if text:
        body(cell, theme, text, size=theme.type_scale.body_sm)
    if items:
        if text:
            spacer(cell, 5)
        bullets(cell, theme, items, colour=ink, size=theme.type_scale.body_sm)


def recommendation_box(container, theme: Theme, *, recommendation: str,
                       rationale: list[str] | None = None,
                       actions: list[str] | None = None,
                       confidence: str = "", title: str = "Recommendation") -> None:
    """The other most-read component. Deliberately visually loudest after the
    cover, because a recommendation a reader skims past is a report that failed."""
    cell = single_cell(container, theme.width, fill=theme.accent_pale,
                       border=(6, theme.accent), accent_bar=(26, theme.accent),
                       pad=theme.card_pad)
    p = para(cell, "", before=0, after=5, keep_with_next=True)
    write(p, title.upper(), font=theme.body, size=theme.type_scale.label, bold=True,
          caps=True, tracking=theme.family.label_tracking, colour=theme.accent)
    if confidence:
        write(p, f"     CONFIDENCE: {confidence.upper()}", font=theme.body,
              size=theme.type_scale.micro, bold=True, caps=True, tracking=1.2,
              colour=theme.palette.ink_soft)
    para(cell, recommendation, font=theme.display, size=theme.type_scale.h3 + 0.5,
         bold=True, colour=theme.primary, before=0, after=6, line=1.24)
    if rationale:
        prose(cell, theme, rationale, size=theme.type_scale.body_sm, gap=4)
    if actions:
        spacer(cell, 7)
        eyebrow(cell, theme, "Next steps", colour=theme.palette.ink_soft, after=4)
        numbered(cell, theme, actions, gap=4)


def risk_box(container, theme: Theme, *, title: str = "Risks and mitigations",
             risks: list[tuple[str, str, str]]) -> None:
    """``risks`` is ``(risk, severity, mitigation)``. Severity drives the tone
    chip, so severity is never carried by colour alone — it is also a word."""
    eyebrow(container, theme, title, colour=theme.palette.alert)
    widths = [theme.width * 0.40, theme.width * 0.14, theme.width * 0.46]
    table = new_table(container, len(risks) + 1, 3, widths)
    table_caption(table, title, "Columns: Risk, Severity, Mitigation")
    head = table.rows[0]
    repeat_header(head)
    for index, label_text in enumerate(("Risk", "Severity", "Mitigation")):
        cell = clear(head.cells[index])
        shade(cell, theme.table_head)
        cell_margins(cell, *theme.cell_pad)
        cell_borders(cell, bottom=(8, theme.accent), top=None, left=None, right=None)
        para(cell, label_text, font=theme.body, size=theme.type_scale.table_head,
             bold=True, caps=True, tracking=1.1, colour=theme.on_table_head,
             before=0, after=0)
    for r_index, (risk, severity, mitigation) in enumerate(risks):
        row = table.rows[r_index + 1]
        keep_row_together(row)
        fill, line, ink, glyph = _tone(theme, severity)
        for c_index, value in enumerate((risk, f"{glyph}  {severity.upper()}", mitigation)):
            cell = clear(row.cells[c_index])
            shade(cell, fill if c_index == 1 else theme.palette.paper)
            cell_margins(cell, *theme.cell_pad)
            valign(cell, "top")
            cell_borders(cell, bottom=(4, theme.palette.line), top=None,
                         left=None, right=None)
            para(cell, value, font=theme.body, size=theme.type_scale.table,
                 bold=c_index == 1, colour=ink if c_index == 1 else theme.palette.ink,
                 before=0, after=0, line=1.24)


def metric_panel(container, theme: Theme, metrics: list[tuple[str, str, str]],
                 *, columns: int | None = None) -> None:
    """Key-metric row. ``metrics`` is ``(label, value, note)``.

    Two rows of cells rather than one so a long value wraps under its caption
    instead of shrinking it. Caps at four across; more than four stacks into a
    second panel, because five metrics at A4 width are illegible.
    """
    columns = columns or min(len(metrics), 4)
    chunks = [metrics[i:i + columns] for i in range(0, len(metrics), columns)]
    for chunk in chunks:
        width = theme.width / len(chunk)
        table = new_table(container, 2, len(chunk), [width] * len(chunk))
        for index, (label_text, value, note) in enumerate(chunk):
            head = clear(table.rows[0].cells[index])
            shade(head, theme.primary)
            cell_margins(head, 80, 120, 60, 120)
            cell_borders(head, right=(24, theme.palette.paper) if index < len(chunk) - 1 else None)
            para(head, label_text, font=theme.body, size=theme.type_scale.micro,
                 bold=True, caps=True, tracking=theme.family.label_tracking,
                 colour=theme.accent, before=0, after=0, keep_with_next=True)

            cell = clear(table.rows[1].cells[index])
            shade(cell, theme.accent_pale)
            cell_margins(cell, 110, 120, 110, 120)
            cell_borders(cell, bottom=(8, theme.accent),
                         right=(24, theme.palette.paper) if index < len(chunk) - 1 else None)
            para(cell, value, font=theme.numeric, size=theme.type_scale.metric_lg,
                 bold=True, colour=theme.primary, before=0, after=0, line=1.05)
            if note:
                para(cell, note, font=theme.body, size=theme.type_scale.micro,
                     colour=theme.palette.ink_soft, before=2, after=0, line=1.2)
        if chunk is not chunks[-1]:
            spacer(container, 4)


def info_card(container, theme: Theme, *, title: str,
              fields: list[tuple[str, str]], columns: int = 2,
              footnote: str = "") -> None:
    """Client, organisation, property or finance summary card. One component
    covers all four because they are the same shape: a titled block of
    label/value pairs that must survive any value length."""
    cell = _card_frame(container, theme, tone_fill=theme.palette.mist,
                       tone_line=theme.palette.line, accent=theme.accent)
    eyebrow(cell, theme, title)
    definition_grid(cell, theme, fields, columns=columns,
                    width=inner_width(theme))
    if footnote:
        para(cell, footnote, font=theme.body, size=theme.type_scale.micro,
             italic=True, colour=theme.palette.ink_faint, before=6, after=0)


def definition_grid(container, theme: Theme, fields: list[tuple[str, str]],
                    *, columns: int = 2, width: float | None = None,
                    input_style: bool = False) -> None:
    """Label/value pairs on a fixed grid.

    ``input_style`` switches values to the field affordance — pale fill and a
    coloured underline — which is what makes a template completable on screen.
    Values live in table cells, so Tab walks the form natively.
    """
    width = width or theme.width
    per = width / columns
    label_w = per * (0.40 if columns > 1 else 0.24)
    widths: list[float] = []
    for _ in range(columns):
        widths += [label_w, per - label_w]

    rows = (len(fields) + columns - 1) // columns
    table = new_table(container, rows, columns * 2, widths)
    for index, (label_text, value) in enumerate(fields):
        r, c = divmod(index, columns)
        row = table.rows[r]
        row_height(row, 7.6)
        label_cell = clear(row.cells[c * 2])
        shade(label_cell, theme.palette.cloud)
        cell_margins(label_cell, *theme.cell_pad)
        valign(label_cell, "center")
        cell_borders(label_cell, top=(4, theme.palette.paper),
                     bottom=(4, theme.palette.paper), left=None, right=None)
        para(label_cell, label_text, font=theme.body, size=theme.type_scale.label,
             bold=True, caps=True, tracking=0.8, colour=theme.palette.ink,
             before=0, after=0, line=1.12)

        value_cell = clear(row.cells[c * 2 + 1])
        shade(value_cell, theme.palette.field if input_style else theme.palette.paper)
        cell_margins(value_cell, *theme.cell_pad)
        valign(value_cell, "center")
        cell_borders(
            value_cell,
            top=(4, theme.palette.paper),
            bottom=(6, theme.palette.field_line) if input_style else (4, theme.palette.line),
            left=None, right=None)
        para(value_cell, value, font=theme.body, size=theme.type_scale.body_sm,
             colour=_placeholder_colour(theme, value), italic=_is_token(value),
             before=0, after=0, line=1.16)

    for index in range(len(fields), rows * columns):
        r, c = divmod(index, columns)
        for offset in (0, 1):
            blank = clear(table.rows[r].cells[c * 2 + offset])
            shade(blank, theme.palette.paper)
            cell_margins(blank, *theme.cell_pad)
            para(blank, "", before=0, after=0)


def responsibility_columns(container, theme: Theme, left: tuple[str, list[str]],
                           right: tuple[str, list[str]],
                           tones: tuple[str, str] = ("brand", "gold")) -> None:
    """Two headed lists side by side — strengths and concerns, included and
    excluded, may and must not.

    A comparison a reader takes in without reading. The two columns are always
    the same width, so neither side reads as the more important one, and the
    tone pair carries the polarity without either column needing the word
    "not" in its first line.
    """
    half = theme.width / 2
    table = new_table(container, 2, 2, [half, half])
    keep_row_together(table.rows[0])
    p = theme.palette
    palettes = {
        "brand": (theme.accent, theme.accent_tint),
        "gold": (theme.primary, p.cloud),
        "success": (p.success, p.success_soft),
        "alert": (p.alert, p.alert_soft),
        "info": (p.info, p.info_soft),
        "warning": (p.warning, p.warning_soft),
    }

    for index, ((title, items), tone) in enumerate(zip((left, right), tones)):
        accent, tint = palettes.get(tone, palettes["brand"])
        head = clear(table.rows[0].cells[index])
        shade(head, accent)
        cell_margins(head, *theme.cell_pad)
        cell_borders(head, top=None, left=None, bottom=None,
                     right=(30, p.paper) if index == 0 else None)
        para(head, title, font=theme.body, size=theme.type_scale.label + 0.5, bold=True,
             caps=True, tracking=theme.family.label_tracking, colour=p.ink_invert,
             before=0, after=0, line=1.15, keep_with_next=True)

        cell = clear(table.rows[1].cells[index])
        shade(cell, tint)
        cell_margins(cell, theme.cell_pad[0] + 40, theme.cell_pad[1])
        cell_borders(cell, top=None, left=None, bottom=(4, p.line),
                     right=(30, p.paper) if index == 0 else None)
        for item_index, item in enumerate(items):
            para_p = para(cell, "", before=0 if item_index == 0 else 4, after=0,
                          line=theme.family.body_line, left_indent=10, hanging=10)
            write(para_p, f"{BULLET}  ", font=theme.body,
                  size=theme.type_scale.body_sm, bold=True, colour=accent)
            write(para_p, item, font=theme.body, size=theme.type_scale.body_sm,
                  colour=p.ink)


def adviser_profile(container, theme: Theme, *, bio: str = "",
                    credentials: list[str] | None = None) -> None:
    photo_w = 30.0
    table = new_table(container, 1, 2, [photo_w, theme.width - photo_w])
    photo = clear(table.rows[0].cells[0])
    shade(photo, theme.palette.cloud)
    cell_margins(photo, 400, 90, 400, 90)
    cell_borders(photo, top=(4, theme.palette.line_strong),
                 left=(4, theme.palette.line_strong),
                 bottom=(4, theme.palette.line_strong),
                 right=(4, theme.palette.line_strong))
    para(photo, "PHOTO", font=theme.body, size=theme.type_scale.micro, bold=True,
         caps=True, tracking=1.2, colour=theme.palette.ink_faint,
         align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=0)

    detail = clear(table.rows[0].cells[1])
    shade(detail, theme.palette.mist)
    cell_margins(detail, *theme.card_pad)
    cell_borders(detail, top=(4, theme.palette.line), bottom=(4, theme.palette.line),
                 right=(4, theme.palette.line), left=None)
    eyebrow(detail, theme, "Your adviser")
    para(detail, theme.brand.author_name, font=theme.display,
         size=theme.type_scale.h3, bold=True, colour=theme.primary,
         before=0, after=1)
    para(detail, f"{theme.brand.author_title}   {ARROW}   {theme.brand.author_credentials}",
         font=theme.body, size=theme.type_scale.body_sm, colour=theme.palette.ink_soft,
         before=0, after=5)
    if bio:
        body(detail, theme, bio, size=theme.type_scale.body_sm)
    if credentials:
        spacer(detail, 5)
        bullets(detail, theme, credentials, size=theme.type_scale.body_sm)
    para(detail, f"{theme.brand.author_email}   ·   {theme.brand.author_phone}",
         font=theme.body, size=theme.type_scale.micro, colour=theme.palette.ink_faint,
         before=6, after=0)


# ==========================================================================
# Tables
# ==========================================================================

def data_table(container, theme: Theme, headers: list[str],
               rows: list[list[str]], widths: list[float] | None = None,
               *, aligns: list[str] | None = None, numeric_cols: set[int] | None = None,
               emphasis_rows: set[int] | None = None,
               total_row: list[str] | None = None,
               caption: str = "", note: str = "") -> None:
    """The library's one table. Comparison, financial, scenario and
    property-feature tables are all this component with different columns —
    which is why they stay visually identical across 36 templates.

    Header rows repeat across pages, rows never split, and columns are pinned
    to millimetre widths derived from the content width so nothing can escape
    the margin.
    """
    p, ts, fam = theme.palette, theme.type_scale, theme.family
    if widths is None:
        widths = [theme.width / len(headers)] * len(headers)
    else:
        scale = theme.width / sum(widths)
        widths = [w * scale for w in widths]
    numeric_cols = numeric_cols or set()
    emphasis_rows = emphasis_rows or set()

    if caption:
        eyebrow(container, theme, caption, colour=p.ink_soft, after=4)

    body_rows = list(rows) + ([total_row] if total_row else [])
    table = new_table(container, len(body_rows) + 1, len(headers), widths)
    table_caption(table, caption or "Data table",
                  f"Columns: {', '.join(headers)}")

    head = table.rows[0]
    repeat_header(head)
    row_height(head, 7.4)
    for index, text in enumerate(headers):
        cell = clear(head.cells[index])
        shade(cell, theme.table_head)
        cell_margins(cell, *theme.cell_pad)
        valign(cell, "bottom")
        cell_borders(cell, bottom=(8, theme.accent), top=None,
                     left=None if fam.table_style != "boxed" else (4, theme.accent),
                     right=None if fam.table_style != "boxed" else (4, theme.accent))
        para(cell, text, font=theme.body, size=ts.table_head, bold=True, caps=True,
             tracking=1.0, colour=theme.on_table_head, before=0, after=0, line=1.12,
             align=WD_ALIGN_PARAGRAPH.RIGHT if index in numeric_cols else None)

    for r_index, values in enumerate(body_rows):
        row = table.rows[r_index + 1]
        keep_row_together(row)
        row_height(row, 7.0)
        is_total = total_row is not None and r_index == len(body_rows) - 1
        is_emphasis = r_index in emphasis_rows

        if is_total:
            tint = theme.accent_tint
        elif is_emphasis:
            tint = p.mist
        elif fam.table_zebra and r_index % 2 == 1:
            tint = p.mist
        else:
            tint = p.paper

        for c_index, value in enumerate(values):
            cell = clear(row.cells[c_index])
            shade(cell, tint)
            cell_margins(cell, *theme.cell_pad)
            valign(cell, "top")
            edges = {"top": None, "left": None, "right": None}
            if fam.table_style == "boxed":
                edges = {"top": (4, p.line), "left": (4, p.line), "right": (4, p.line)}
            elif fam.table_style == "ledger" and c_index > 0:
                edges["left"] = (4, p.line)
            cell_borders(cell,
                         bottom=(8, theme.accent) if is_total else (4, p.line),
                         **edges)
            numeric = c_index in numeric_cols
            para(cell, value,
                 font=theme.numeric if numeric else theme.body,
                 size=ts.table,
                 bold=is_total or (is_emphasis and c_index == 0),
                 colour=theme.primary if is_total else _placeholder_colour(theme, value),
                 italic=_is_token(value) and not is_total,
                 before=0, after=0, line=1.22,
                 align=WD_ALIGN_PARAGRAPH.RIGHT if numeric else None)

    if note:
        para(container, note, font=theme.body, size=ts.micro, italic=True,
             colour=p.ink_faint, before=4, after=0, line=1.24)


def comparison_table(container, theme: Theme, *, subject_labels: list[str],
                     attributes: list[tuple[str, list[str]]],
                     caption: str = "", winner_index: int | None = None) -> None:
    """Side-by-side comparison — properties, loans, scenarios, lenders.

    Attribute-per-row rather than subject-per-row, because a comparison is read
    across a single attribute at a time, and because adding a fourth subject
    should widen the table rather than force a page break.
    """
    label_w = theme.width * 0.26
    col_w = (theme.width - label_w) / len(subject_labels)
    headers = [""] + subject_labels
    widths = [label_w] + [col_w] * len(subject_labels)

    if caption:
        eyebrow(container, theme, caption, colour=theme.palette.ink_soft, after=4)

    table = new_table(container, len(attributes) + 1, len(headers), widths)
    table_caption(table, caption or "Comparison",
                  f"Compares: {', '.join(subject_labels)}")
    head = table.rows[0]
    repeat_header(head)
    row_height(head, 9.0)
    for index, text in enumerate(headers):
        cell = clear(head.cells[index])
        highlight = winner_index is not None and index == winner_index + 1
        shade(cell, theme.accent if highlight else theme.table_head)
        cell_margins(cell, *theme.cell_pad)
        valign(cell, "center")
        cell_borders(cell, bottom=(8, theme.accent),
                     right=(16, theme.palette.paper) if index < len(headers) - 1 else None,
                     top=None, left=None)
        para(cell, text, font=theme.body, size=theme.type_scale.table_head + 0.5,
             bold=True, caps=True, tracking=1.0, colour=theme.on_table_head,
             before=0, after=0, line=1.14,
             align=WD_ALIGN_PARAGRAPH.CENTER if index else None)

    for r_index, (attribute, values) in enumerate(attributes):
        row = table.rows[r_index + 1]
        keep_row_together(row)
        tint = theme.palette.mist if theme.family.table_zebra and r_index % 2 == 1 \
            else theme.palette.paper
        label_cell = clear(row.cells[0])
        shade(label_cell, theme.palette.cloud)
        cell_margins(label_cell, *theme.cell_pad)
        valign(label_cell, "center")
        cell_borders(label_cell, bottom=(4, theme.palette.line), top=None,
                     left=None, right=None)
        para(label_cell, attribute, font=theme.body, size=theme.type_scale.label,
             bold=True, caps=True, tracking=0.8, colour=theme.palette.ink,
             before=0, after=0, line=1.14)
        for c_index, value in enumerate(values):
            cell = clear(row.cells[c_index + 1])
            highlight = winner_index is not None and c_index == winner_index
            shade(cell, theme.accent_pale if highlight else tint)
            cell_margins(cell, *theme.cell_pad)
            valign(cell, "center")
            cell_borders(cell, bottom=(4, theme.palette.line), top=None,
                         left=None, right=None)
            para(cell, value, font=theme.body, size=theme.type_scale.table,
                 colour=_placeholder_colour(theme, value), italic=_is_token(value),
                 before=0, after=0, line=1.2, align=WD_ALIGN_PARAGRAPH.CENTER)


def status_table(container, theme: Theme, *, headers: list[str],
                 rows: list[tuple[list[str], str]], widths: list[float] | None = None,
                 status_col: int = -1, caption: str = "", note: str = "") -> None:
    """Compliance/verification table where one column is a status chip.

    ``rows`` is ``(values, status_keyword)``. The status cell renders as a
    filled chip with a glyph *and* the word, so the state survives grayscale
    printing and colour-blind readers.
    """
    if widths is None:
        widths = [theme.width / len(headers)] * len(headers)
    else:
        scale = theme.width / sum(widths)
        widths = [w * scale for w in widths]
    if caption:
        eyebrow(container, theme, caption, colour=theme.palette.ink_soft, after=4)

    table = new_table(container, len(rows) + 1, len(headers), widths)
    table_caption(table, caption or "Status table",
                  f"Columns: {', '.join(headers)}")
    head = table.rows[0]
    repeat_header(head)
    for index, text in enumerate(headers):
        cell = clear(head.cells[index])
        shade(cell, theme.table_head)
        cell_margins(cell, *theme.cell_pad)
        cell_borders(cell, bottom=(8, theme.accent), top=None, left=None, right=None)
        para(cell, text, font=theme.body, size=theme.type_scale.table_head, bold=True,
             caps=True, tracking=1.0, colour=theme.on_table_head, before=0, after=0,
             line=1.12)

    normalised = status_col if status_col >= 0 else len(headers) + status_col
    for r_index, (values, status) in enumerate(rows):
        row = table.rows[r_index + 1]
        keep_row_together(row)
        fill, line, ink, glyph = _tone(theme, status)
        for c_index, value in enumerate(values):
            cell = clear(row.cells[c_index])
            is_status = c_index == normalised
            shade(cell, fill if is_status else theme.palette.paper)
            cell_margins(cell, *theme.cell_pad)
            valign(cell, "top")
            cell_borders(cell, bottom=(4, theme.palette.line), top=None,
                         left=None, right=None)
            para(cell, f"{glyph}  {value}" if is_status else value,
                 font=theme.body, size=theme.type_scale.table, bold=is_status,
                 colour=ink if is_status else _placeholder_colour(theme, value),
                 italic=_is_token(value) and not is_status,
                 before=0, after=0, line=1.22)

    if note:
        para(container, note, font=theme.body, size=theme.type_scale.micro, italic=True,
             colour=theme.palette.ink_faint, before=4, after=0, line=1.24)


def bar_chart(container, theme: Theme, *, rows: list[tuple[str, float, str]],
              caption: str = "", note: str = "", maximum: float | None = None) -> None:
    """Native horizontal bar chart drawn with shaded table cells.

    No image dependency, so it survives Word round-trips and grayscale printing,
    and the value is always printed as text beside the bar. Use for anything
    ranked or proportional — lender comparison, expense split, scenario
    outcomes. Use ``chart_frame`` instead where a real chart image is generated.
    """
    if caption:
        eyebrow(container, theme, caption, colour=theme.palette.ink_soft, after=4)
    peak = maximum or max((value for _, value, _ in rows), default=1) or 1
    label_w = theme.width * 0.28
    value_w = theme.width * 0.16
    track_w = theme.width - label_w - value_w
    series = theme.palette.series

    table = new_table(container, len(rows), 3, [label_w, track_w, value_w])
    for index, (label_text, value, display) in enumerate(rows):
        row = table.rows[index]
        row_height(row, 7.4)
        keep_row_together(row)

        label_cell = clear(row.cells[0])
        cell_margins(label_cell, *theme.cell_pad)
        valign(label_cell, "center")
        cell_borders(label_cell, bottom=(4, theme.palette.line), top=None,
                     left=None, right=None)
        para(label_cell, label_text, font=theme.body, size=theme.type_scale.body_sm,
             colour=theme.palette.ink, before=0, after=0, line=1.16)

        track = clear(row.cells[1])
        cell_margins(track, 40, 0, 40, 0)
        valign(track, "center")
        cell_borders(track, bottom=(4, theme.palette.line), top=None,
                     left=None, right=None)
        ratio = max(0.02, min(1.0, (value / peak) if peak else 0))
        filled = max(1.0, track_w * ratio)
        inner = new_table(track, 1, 2, [filled, max(0.6, track_w - filled)])
        inner.alignment = WD_TABLE_ALIGNMENT.LEFT
        bar_cell = clear(inner.cell(0, 0))
        shade(bar_cell, series[index % len(series)])
        cell_margins(bar_cell, 0, 0, 0, 0)
        bar_para = para(bar_cell, "", before=0, after=0)
        bar_para.paragraph_format.line_spacing = Pt(9)
        bar_para.add_run().font.size = Pt(1)
        rest = clear(inner.cell(0, 1))
        shade(rest, theme.palette.cloud)
        cell_margins(rest, 0, 0, 0, 0)
        rest_para = para(rest, "", before=0, after=0)
        rest_para.paragraph_format.line_spacing = Pt(9)
        rest_para.add_run().font.size = Pt(1)

        value_cell = clear(row.cells[2])
        cell_margins(value_cell, *theme.cell_pad)
        valign(value_cell, "center")
        cell_borders(value_cell, bottom=(4, theme.palette.line), top=None,
                     left=None, right=None)
        para(value_cell, display, font=theme.numeric, size=theme.type_scale.body_sm,
             bold=True, colour=theme.primary, before=0, after=0,
             align=WD_ALIGN_PARAGRAPH.RIGHT)

    if note:
        para(container, note, font=theme.body, size=theme.type_scale.micro,
             italic=True, colour=theme.palette.ink_faint, before=4, after=0)


def chart_frame(container, theme: Theme, *, title: str, kind: str,
                height_mm: float = 62.0, caption: str = "", source: str = "",
                alt_text: str = "", binding: str = "") -> None:
    """Reserved frame for a chart image rendered by the platform.

    Carries everything the generator and the accessibility reviewer need: the
    chart kind, the binding that fills it, the alt text, and the data source
    line. An empty frame in the master template is a specification, not a gap.
    """
    p, ts = theme.palette, theme.type_scale
    table = new_table(container, 1, 1, [theme.width])
    row_height(table.rows[0], height_mm)
    cell = clear(table.cell(0, 0))
    shade(cell, p.mist if theme.family.chart_style != "outline" else p.paper)
    cell_margins(cell, 150, 180, 150, 180)
    cell_borders(cell, top=(4, p.line_strong), left=(4, p.line_strong),
                 bottom=(4, p.line_strong), right=(4, p.line_strong))
    valign(cell, "center")
    para(cell, title, font=theme.display, size=ts.h4, bold=True, colour=theme.primary,
         align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=3)
    para(cell, f"[  {kind.upper()}  ]", font=theme.body, size=ts.label, bold=True,
         caps=True, tracking=1.8, colour=theme.accent,
         align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=3)
    if binding:
        para(cell, binding, font=theme.numeric, size=ts.micro, colour=p.ink_faint,
             align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=0)
    meta = []
    if caption:
        meta.append(caption)
    if source:
        meta.append(f"Source: {source}")
    if alt_text:
        meta.append(f"Alt text: {alt_text}")
    if meta:
        cell.paragraphs[-1].paragraph_format.keep_with_next = True
        para(container, "   ·   ".join(meta), font=theme.body, size=ts.micro,
             italic=True, colour=p.ink_faint, before=4, after=0, line=1.24)


# ==========================================================================
# Imagery
# ==========================================================================

def image_frame(container, theme: Theme, *, caption: str = "", height_mm: float = 70.0,
                placeholder: str = "", alt_text: str = "",
                width_mm: float | None = None) -> None:
    p, ts = theme.palette, theme.type_scale
    width = width_mm or theme.width
    table = new_table(container, 1, 1, [width])
    row_height(table.rows[0], height_mm)
    cell = clear(table.cell(0, 0))
    shade(cell, p.cloud)
    cell_margins(cell, 120, 140, 120, 140)
    if theme.family.image_style == "framed":
        cell_borders(cell, top=(6, p.line_strong), left=(6, p.line_strong),
                     bottom=(6, p.line_strong), right=(6, p.line_strong))
    valign(cell, "center")
    para(cell, placeholder or "[  IMAGE  ]", font=theme.body, size=ts.label,
         bold=True, caps=True, tracking=1.6, colour=p.ink_faint,
         align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=0)
    lines = [t for t in (caption, f"Alt text: {alt_text}" if alt_text else "") if t]
    if lines:
        cell.paragraphs[-1].paragraph_format.keep_with_next = True
        para(container, "   ·   ".join(lines), font=theme.body, size=ts.micro,
             italic=True, colour=p.ink_faint, before=3, after=0, line=1.22)


def image_gallery(container, theme: Theme, *, count: int = 4, columns: int = 2,
                  height_mm: float = 44.0, captions: list[str] | None = None) -> None:
    captions = captions or [f"{{{{property.images.{i}.caption}}}}" for i in range(count)]
    rows = (count + columns - 1) // columns
    gutter = 3.0
    width = (theme.width - gutter * (columns - 1)) / columns
    for r in range(rows):
        widths: list[float] = []
        for c in range(columns):
            widths.append(width)
            if c < columns - 1:
                widths.append(gutter)
        table = new_table(container, 1, len(widths), widths)
        for c in range(columns):
            index = r * columns + c
            cell = clear(table.rows[0].cells[c * 2])
            if index >= count:
                continue
            shade(cell, theme.palette.cloud)
            cell_margins(cell, 100, 100, 100, 100)
            cell_borders(cell, top=(4, theme.palette.line_strong),
                         left=(4, theme.palette.line_strong),
                         bottom=(4, theme.palette.line_strong),
                         right=(4, theme.palette.line_strong))
            row_height(table.rows[0], height_mm)
            valign(cell, "center")
            para(cell, f"[  IMAGE {index + 1}  ]", font=theme.body,
                 size=theme.type_scale.micro, bold=True, caps=True, tracking=1.4,
                 colour=theme.palette.ink_faint, align=WD_ALIGN_PARAGRAPH.CENTER,
                 before=0, after=0)
            para(cell, captions[index] if index < len(captions) else "",
                 font=theme.body, size=theme.type_scale.micro, italic=True,
                 colour=theme.palette.ink_faint, align=WD_ALIGN_PARAGRAPH.CENTER,
                 before=3, after=0)
        if r < rows - 1:
            spacer(container, 3)


def map_frame(container, theme: Theme, *, title: str = "Location",
              height_mm: float = 76.0, legend: list[tuple[str, str]] | None = None,
              alt_text: str = "") -> None:
    image_frame(container, theme, height_mm=height_mm,
                placeholder=f"[  {title.upper()} MAP  ]", alt_text=alt_text)
    if legend:
        spacer(container, 3)
        width = theme.width / len(legend)
        table = new_table(container, 1, len(legend), [width] * len(legend))
        for index, (marker, label_text) in enumerate(legend):
            cell = clear(table.rows[0].cells[index])
            cell_margins(cell, 50, 60, 50, 60)
            p = para(cell, "", before=0, after=0)
            write(p, f"{marker}  ", font=theme.body, size=theme.type_scale.body_sm,
                  bold=True, colour=theme.palette.series[index % 10])
            write(p, label_text, font=theme.body, size=theme.type_scale.micro,
                  colour=theme.palette.ink_soft)


# ==========================================================================
# Process, timeline, checklist
# ==========================================================================

def timeline(container, theme: Theme, entries: list[tuple[str, str, str]],
             *, caption: str = "") -> None:
    """``entries`` is ``(when, what, detail)``. Vertical, because a horizontal
    timeline cannot survive a fifth entry at A4 width."""
    if caption:
        eyebrow(container, theme, caption, colour=theme.palette.ink_soft, after=4)
    when_w = theme.width * 0.20
    marker_w = 8.0
    table = new_table(container, len(entries), 3,
                      [when_w, marker_w, theme.width - when_w - marker_w])
    for index, (when, what, detail) in enumerate(entries):
        row = table.rows[index]
        keep_row_together(row)
        row_height(row, 10.0)

        when_cell = clear(row.cells[0])
        cell_margins(when_cell, *theme.cell_pad)
        valign(when_cell, "top")
        para(when_cell, when, font=theme.body, size=theme.type_scale.label, bold=True,
             caps=True, tracking=1.0, colour=theme.accent, before=0, after=0,
             align=WD_ALIGN_PARAGRAPH.RIGHT, line=1.2)

        marker = clear(row.cells[1])
        shade(marker, theme.accent_tint)
        cell_margins(marker, 0, 0, 0, 0)
        cell_borders(marker, left=(14, theme.accent), top=None, bottom=None, right=None)
        para(marker, "", before=0, after=0)

        detail_cell = clear(row.cells[2])
        cell_margins(detail_cell, *theme.cell_pad)
        valign(detail_cell, "top")
        cell_borders(detail_cell, bottom=(4, theme.palette.line), top=None,
                     left=None, right=None)
        para(detail_cell, what, font=theme.body, size=theme.type_scale.body,
             bold=True, colour=theme.primary, before=0, after=1, line=1.2)
        if detail:
            para(detail_cell, detail, font=theme.body, size=theme.type_scale.body_sm,
                 colour=theme.palette.ink_soft, before=0, after=0, line=1.22)


def process_flow(container, theme: Theme, steps: list[tuple[str, str]],
                 *, caption: str = "") -> None:
    """Numbered stage ladder — referral workflows, acquisition process,
    onboarding. ``steps`` is ``(name, description)``."""
    if caption:
        eyebrow(container, theme, caption, colour=theme.palette.ink_soft, after=4)
    num_w, name_w = 12.0, 40.0
    table = new_table(container, len(steps), 3,
                      [num_w, name_w, theme.width - num_w - name_w])
    for index, (name, description) in enumerate(steps):
        row = table.rows[index]
        keep_row_together(row)
        row_height(row, 8.6)
        tint = theme.palette.paper if index % 2 == 0 else theme.palette.mist

        badge = clear(row.cells[0])
        shade(badge, theme.primary)
        cell_margins(badge, 80, 40, 80, 40)
        valign(badge, "center")
        cell_borders(badge, top=(4, theme.palette.paper),
                     bottom=(4, theme.palette.paper), left=None, right=None)
        para(badge, str(index + 1), font=theme.display, size=theme.type_scale.body + 1,
             bold=True, colour=theme.accent, align=WD_ALIGN_PARAGRAPH.CENTER,
             before=0, after=0)

        name_cell = clear(row.cells[1])
        shade(name_cell, tint)
        cell_margins(name_cell, *theme.cell_pad)
        valign(name_cell, "center")
        cell_borders(name_cell, top=(4, theme.palette.paper),
                     bottom=(4, theme.palette.paper), left=None, right=None)
        para(name_cell, name, font=theme.body, size=theme.type_scale.label + 0.5,
             bold=True, caps=True, tracking=1.1, colour=theme.primary,
             before=0, after=0, line=1.14)

        desc = clear(row.cells[2])
        shade(desc, tint)
        cell_margins(desc, *theme.cell_pad)
        valign(desc, "center")
        cell_borders(desc, top=(4, theme.palette.paper),
                     bottom=(4, theme.palette.paper), left=None, right=None)
        para(desc, description, font=theme.body, size=theme.type_scale.body_sm,
             colour=theme.palette.ink, before=0, after=0, line=1.2)


def checklist(container, theme: Theme, items: list[str], *, title: str = "",
              columns: int = 1, with_owner: bool = False) -> None:
    if title:
        eyebrow(container, theme, title, colour=theme.palette.ink_soft, after=4)
    if with_owner:
        data_table(
            container, theme,
            ["", "Item", "Owner", "Due", "Status"],
            [[CHECKBOX, item, "{{task.owner}}", "{{task.due}}", ""] for item in items],
            widths=[8, 92, 32, 24, 26],
        )
        return
    per = theme.width / columns
    rows = (len(items) + columns - 1) // columns
    table = new_table(container, rows, columns, [per] * columns)
    for index, item in enumerate(items):
        r, c = divmod(index, columns)
        cell = clear(table.rows[r].cells[c])
        cell_margins(cell, *theme.cell_pad)
        cell_borders(cell, bottom=(4, theme.palette.line), top=None,
                     left=None, right=None)
        p = para(cell, "", before=0, after=0, line=1.2, left_indent=12, hanging=12)
        write(p, f"{CHECKBOX}  ", font=theme.body, size=theme.type_scale.body,
              colour=theme.accent)
        write(p, item, font=theme.body, size=theme.type_scale.body_sm,
              colour=theme.palette.ink)


# ==========================================================================
# Governance and closing matter
# ==========================================================================

def document_control(container, theme: Theme, *, compact: bool = False,
                     extra: list[tuple[str, str]] | None = None) -> None:
    b = theme.brand
    fields = [
        ("Prepared for", b.client_name),
        ("Client reference", b.client_reference),
        ("Prepared by", b.author_name),
        ("Date of issue", b.issue_date),
        ("Document reference", b.document_reference),
        ("Version", b.version),
    ]
    if not compact:
        fields += [
            ("Recipient", b.recipient_name),
            ("Confidentiality", b.confidentiality),
            ("Legal entity", b.legal_entity_name),
            ("ABN", b.abn),
        ]
    fields += extra or []
    if compact:
        definition_grid(container, theme, fields, columns=2)
    else:
        info_card(container, theme, title="Document control", fields=fields, columns=2)


def signature_block(container, theme: Theme,
                    blocks: list[tuple[str, list[str]]]) -> None:
    width = theme.width / len(blocks)
    table = new_table(container, 1, len(blocks), [width] * len(blocks))
    keep_row_together(table.rows[0])
    for index, (title, lines) in enumerate(blocks):
        cell = clear(table.rows[0].cells[index])
        shade(cell, theme.palette.paper)
        cell_margins(cell, 170, 190, 200, 190)
        cell_borders(cell, top=(12, theme.accent), left=(4, theme.palette.line),
                     bottom=(4, theme.palette.line), right=(4, theme.palette.line))
        para(cell, title, font=theme.body, size=theme.type_scale.label + 0.5,
             bold=True, caps=True, tracking=theme.family.label_tracking,
             colour=theme.primary, before=0, after=6, keep_with_next=True)
        for line_index, line in enumerate(lines):
            label_text, _, value = line.partition(":")
            p = para(cell, "", before=0 if line_index == 0 else 7, after=0, line=1.2)
            write(p, f"{label_text.strip()}   ", font=theme.body,
                  size=theme.type_scale.micro, bold=True, caps=True, tracking=1.0,
                  colour=theme.palette.ink_soft)
            value = value.strip()
            if value:
                write(p, value, font=theme.body, size=theme.type_scale.body_sm,
                      colour=_placeholder_colour(theme, value), italic=_is_token(value))
            else:
                write(p, "_" * 26, font=theme.body, size=theme.type_scale.body_sm,
                      colour=theme.palette.line_strong)


def approval_block(container, theme: Theme,
                   approvals: list[tuple[str, str, str]]) -> None:
    """``approvals`` is ``(role, name, status)`` — the governance trail that
    makes a generated document auditable."""
    status_table(
        container, theme,
        headers=["Role", "Name", "Date", "Status"],
        rows=[([role, name, "{{approval.date}}", status], status)
              for role, name, status in approvals],
        widths=[46, 52, 30, 34],
        caption="Approvals",
    )


def appendix_opener(container, theme: Theme, letter: str, title: str,
                    description: str = "") -> None:
    page_break(container)
    para(container, f"APPENDIX {letter}", font=theme.body,
         size=theme.type_scale.label, bold=True, caps=True,
         tracking=theme.family.label_tracking + 1.2, colour=theme.accent,
         before=0, after=4, keep_with_next=True)
    para(container, title, font=theme.display, size=theme.type_scale.h1, bold=True,
         colour=theme.primary, before=0, after=6, keep_with_next=True)
    rule(container, theme, colour=theme.accent, weight=10, after=6)
    if description:
        body(container, theme, description, colour=theme.palette.ink_soft)


def disclaimer_page(container, theme: Theme, *, extra_sections:
                    list[tuple[str, str]] | None = None) -> None:
    page_break(container)
    section_opener(container, theme, "", "Important information",
                   "Disclaimer, privacy and terms")
    gap(container, theme)
    sections = [
        ("Disclaimer", theme.brand.disclaimer),
        ("Privacy", theme.brand.privacy_notice),
        ("Terms of use", theme.brand.terms),
    ] + (extra_sections or [])
    for title, text in sections:
        subsection(container, theme, title)
        body(container, theme, text, size=theme.type_scale.body_sm,
             colour=theme.palette.ink_soft)
        spacer(container, 5)


def contact_page(container, theme: Theme, *, offices:
                 list[tuple[str, str]] | None = None) -> None:
    gap(container, theme)
    cell = single_cell(container, theme.width, fill=theme.primary,
                       border=(4, theme.primary), pad=(320, 260))
    _logo_slot(cell, theme, theme.brand.logo_placeholder, 54.0, on_dark=True)
    spacer(cell, 10)
    para(cell, theme.brand.organisation_name, font=theme.display,
         size=theme.type_scale.h1, bold=True, colour=theme.palette.ink_invert,
         before=0, after=3)
    para(cell, theme.brand.tagline, font=theme.body, size=theme.type_scale.label,
         bold=True, caps=True, tracking=theme.family.label_tracking + 0.6,
         colour=theme.accent, before=0, after=10)
    for label_text, value in (("Web", theme.brand.website), ("Email", theme.brand.email),
                              ("Phone", theme.brand.phone),
                              ("Address", theme.brand.address),
                              ("Social", theme.brand.socials)):
        p = para(cell, "", before=0, after=4, line=1.24)
        write(p, f"{label_text.upper()}   ", font=theme.body,
              size=theme.type_scale.micro, bold=True, caps=True, tracking=1.3,
              colour=theme.accent)
        write(p, value, font=theme.body, size=theme.type_scale.body_sm,
              colour=theme.palette.ink_invert_soft)
    if offices:
        spacer(cell, 8)
        for name, address in offices:
            p = para(cell, "", before=0, after=3, line=1.24)
            write(p, f"{name}   ", font=theme.body, size=theme.type_scale.micro,
                  bold=True, caps=True, tracking=1.2, colour=theme.accent)
            write(p, address, font=theme.body, size=theme.type_scale.micro,
                  colour=theme.palette.ink_invert_soft)
    if theme.brand.powered_by:
        spacer(cell, 10)
        para(cell, theme.brand.powered_by, font=theme.body,
             size=theme.type_scale.micro, caps=True, tracking=1.6,
             colour=theme.palette.ink_invert_soft, before=0, after=0)


def back_cover(container, theme: Theme) -> None:
    page_break(container)
    contact_page(container, theme)
    spacer(container, 8)
    para(container,
         f"{theme.brand.legal_entity_name}   ·   ABN {theme.brand.abn}   ·   "
         f"{theme.brand.confidentiality}",
         font=theme.body, size=theme.type_scale.micro, colour=theme.palette.ink_faint,
         align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=0)


def table_of_contents(container, theme: Theme,
                      entries: list[tuple[str, str]]) -> None:
    """Static contents table.

    A live Word ``TOC`` field is deliberately not used: it renders as "right
    click to update" until a human opens the file, which is unacceptable on a
    document that is generated, converted to PDF and emailed without ever being
    opened in Word. The generator knows the section list, so it writes it.
    """
    section_opener(container, theme, "", "Contents", "")
    gap(container, theme)
    for number, title in entries:
        p = para(container, "", before=0, after=6, line=1.2)
        write(p, f"{number}    ", font=theme.numeric, size=theme.type_scale.body,
              bold=True, colour=theme.accent)
        write(p, title, font=theme.body, size=theme.type_scale.body,
              colour=theme.palette.ink)
        horizontal_rule(p, theme.palette.line, size=4, space=3)


def white_label_panel(container, theme: Theme, *, slots:
                      list[tuple[str, str, str]]) -> None:
    """The in-document control sheet listing every configurable area.

    Present in the Aurixa master build only; the generator strips it for
    partner builds. Its job is to make the white-label surface legible to a
    human who never opens the admin UI.
    """
    section_opener(container, theme, "W", "Brand & customisation panel",
                   "White-label control sheet")
    gap(container, theme)
    highlight_box(
        container, theme, title="Delete before issue",
        text=("Every item below is populated from the organisation's branding "
              "configuration when the platform generates this document. The table is "
              "reproduced here so the configurable surface is visible to anyone "
              "reviewing the master template."),
        tone="warning")
    gap(container, theme, 0.6)
    data_table(container, theme, ["Configurable area", "Binding", "Appears in"],
               [[name, binding, where] for name, binding, where in slots],
               widths=[42, 52, 56])


# ==========================================================================
# Document assembly
# ==========================================================================

def base_document(theme: Theme, doc_title: str):
    """A blank document with page setup, running furniture and base styles."""
    from docx import Document

    doc = Document()
    configure_section(doc.sections[0], theme)
    running_header(doc.sections[0], theme, doc_title)
    running_footer(doc.sections[0], theme)

    normal = doc.styles["Normal"]
    set_style_font(normal, theme.body, theme.type_scale.body, theme.palette.ink)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = theme.family.body_line
    return doc


def set_properties(doc, theme: Theme, *, title: str, subject: str,
                   keywords: str, category: str) -> None:
    props = doc.core_properties
    props.title = title
    props.subject = subject
    props.author = theme.brand.organisation_name
    props.category = category
    props.keywords = keywords
    props.comments = (
        "Generated by the Aurixa Command Center template library "
        "(scripts/aurixa-templates). Edit the builder and regenerate rather than "
        "editing this file by hand."
    )
    # White-label level 4 must leave no Aurixa fingerprint, including in metadata.
    if theme.brand.level < 4:
        props.last_modified_by = "Aurixa Systems"
