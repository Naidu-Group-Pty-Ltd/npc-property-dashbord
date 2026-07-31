"""Low-level WordprocessingML primitives shared by every Aurixa document generator.

``python-docx`` deliberately exposes only a small, safe subset of the OOXML
surface. Everything a designed document needs — cell shading, per-edge borders,
cell padding, character tracking, page fields, gradient-free colour bands — has
to be written as raw XML.

The single non-obvious rule this module enforces: **WordprocessingML property
containers are ordered `xsd:sequence` types.** Appending `<w:caps>` after
`<w:color>` inside an `<w:rPr>` produces well-formed XML that Word and
LibreOffice both refuse to open. Every raw element added here is spliced into
its schema-mandated slot instead of appended.

This module is layout-agnostic and brand-agnostic. Visual decisions live in the
component libraries that build on top of it.
"""

from __future__ import annotations

from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor

__all__ = [
    "el", "splice", "get_or_add", "replace_child",
    "shade", "cell_borders", "cell_margins", "valign", "cell_width",
    "table_borders_none", "fixed_layout", "table_indent", "table_cell_spacing",
    "table_caption",
    "row_height", "keep_row_together", "repeat_header", "merge_across",
    "para_format", "write", "para", "clear", "tab_stop",
    "page_field", "hyperlink", "page_break", "spacer", "horizontal_rule",
    "new_table", "single_cell", "set_style_font", "bookmark",
    "WD_ALIGN_PARAGRAPH", "WD_TABLE_ALIGNMENT", "Mm", "Pt", "RGBColor",
]


# ==========================================================================
# Schema-ordered XML helpers
# ==========================================================================

def el(tag: str, **attrs) -> OxmlElement:
    """Build a `w:`-namespaced element with `w:`-namespaced attributes."""
    node = OxmlElement(tag)
    for key, value in attrs.items():
        node.set(qn(f"w:{key}"), str(value))
    return node


#: Ordered child sequences for the property containers we write into.
#: Taken from ECMA-376 Part 1 (CT_RPr, CT_PPr, CT_TcPr, CT_TblPr, CT_TrPr,
#: CT_TblBorders, CT_TcBorders, CT_PBdr, CT_SectPr).
CHILD_ORDER: dict[str, tuple[str, ...]] = {
    "rPr": (
        "rStyle", "rFonts", "b", "bCs", "i", "iCs", "caps", "smallCaps", "strike",
        "dstrike", "outline", "shadow", "emboss", "imprint", "noProof", "snapToGrid",
        "vanish", "webHidden", "color", "spacing", "w", "kern", "position", "sz",
        "szCs", "highlight", "u", "effect", "bdr", "shd", "fitText", "vertAlign",
        "rtl", "cs", "em", "lang", "eastAsianLayout", "specVanish", "oMath",
    ),
    "pPr": (
        "pStyle", "keepNext", "keepLines", "pageBreakBefore", "framePr",
        "widowControl", "numPr", "suppressLineNumbers", "pBdr", "shd", "tabs",
        "suppressAutoHyphens", "kinsoku", "wordWrap", "overflowPunct",
        "topLinePunct", "autoSpaceDE", "autoSpaceDN", "bidi", "adjustRightInd",
        "snapToGrid", "spacing", "ind", "contextualSpacing", "mirrorIndents",
        "suppressOverlap", "jc", "textDirection", "textAlignment", "textboxTightWrap",
        "outlineLvl", "divId", "cnfStyle", "rPr", "sectPr", "pPrChange",
    ),
    "pBdr": ("top", "left", "bottom", "right", "between", "bar"),
    "tcPr": (
        "cnfStyle", "tcW", "gridSpan", "hMerge", "vMerge", "tcBorders", "shd",
        "noWrap", "tcMar", "textDirection", "tcFitText", "vAlign", "hideMark",
    ),
    "tcBorders": ("top", "start", "left", "bottom", "end", "right", "insideH",
                  "insideV", "tl2br", "tr2bl"),
    "tblBorders": ("top", "start", "left", "bottom", "end", "right", "insideH",
                   "insideV"),
    "tblPr": (
        "tblStyle", "tblpPr", "tblOverlap", "bidiVisual", "tblStyleRowBandSize",
        "tblStyleColBandSize", "tblW", "jc", "tblCellSpacing", "tblInd",
        "tblBorders", "shd", "tblLayout", "tblCellMar", "tblLook", "tblCaption",
        "tblDescription",
    ),
    "trPr": (
        "cnfStyle", "divId", "gridBefore", "gridAfter", "wBefore", "wAfter",
        "cantSplit", "trHeight", "tblHeader", "tblCellSpacing", "jc", "hidden",
    ),
}


def splice(parent, node) -> None:
    """Insert ``node`` into ``parent`` at its schema-mandated position."""
    order = CHILD_ORDER.get(parent.tag.split("}")[-1])
    if order is None:
        parent.append(node)
        return
    name = node.tag.split("}")[-1]
    try:
        rank = order.index(name)
    except ValueError:
        parent.append(node)
        return
    for existing in parent:
        existing_name = existing.tag.split("}")[-1]
        if existing_name not in order or order.index(existing_name) > rank:
            existing.addprevious(node)
            return
    parent.append(node)


def get_or_add(parent, tag: str):
    node = parent.find(qn(tag))
    if node is None:
        node = OxmlElement(tag)
        splice(parent, node)
    return node


def replace_child(parent, tag: str, node) -> None:
    """Drop existing ``tag`` children, then splice ``node`` into place."""
    for old in parent.findall(qn(tag)):
        parent.remove(old)
    splice(parent, node)


# ==========================================================================
# Cell and table properties
# ==========================================================================

def shade(cell, fill: str) -> None:
    """Solid background fill on a table cell. ``fill`` is RGB hex, no ``#``."""
    replace_child(cell._tc.get_or_add_tcPr(), "w:shd",
                  el("w:shd", val="clear", color="auto", fill=fill))


def cell_borders(cell, **edges) -> None:
    """Per-edge borders.

    Keywords are ``top`` / ``left`` / ``bottom`` / ``right`` / ``insideH`` /
    ``insideV``. Each takes ``None`` (explicitly no border) or a
    ``(size_eighths_of_a_point, colour_hex)`` tuple. A third element may supply
    the line style, e.g. ``(8, "0C2340", "dashed")``. Omitted edges are left
    untouched, so borders can be layered by successive calls.
    """
    borders = get_or_add(cell._tc.get_or_add_tcPr(), "w:tcBorders")
    for edge, spec in edges.items():
        if spec is None:
            node = el(f"w:{edge}", val="nil")
        else:
            size, colour = spec[0], spec[1]
            style = spec[2] if len(spec) > 2 else "single"
            node = el(f"w:{edge}", val=style, sz=size, space=0, color=colour)
        replace_child(borders, f"w:{edge}", node)


def cell_margins(cell, top: int, left: int, bottom: int | None = None,
                 right: int | None = None) -> None:
    """Cell padding in dxa (twentieths of a point). 20 dxa = 1pt.

    ``bottom`` and ``right`` default to ``top`` and ``left``, so the common
    symmetric case is ``cell_margins(cell, *theme.cell_pad)``.
    """
    if bottom is None:
        bottom = top
    if right is None:
        right = left
    mar = OxmlElement("w:tcMar")
    for tag, value in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        mar.append(el(f"w:{tag}", w=value, type="dxa"))
    replace_child(cell._tc.get_or_add_tcPr(), "w:tcMar", mar)


def valign(cell, value: str = "center") -> None:
    replace_child(cell._tc.get_or_add_tcPr(), "w:vAlign", el("w:vAlign", val=value))


def cell_width(cell, mm: float) -> None:
    replace_child(cell._tc.get_or_add_tcPr(), "w:tcW",
                  el("w:tcW", w=int(Mm(mm).twips), type="dxa"))


def merge_across(row, first: int, last: int):
    """Horizontally merge cells ``first``..``last`` and return the survivor."""
    survivor = row.cells[first]
    for index in range(first + 1, last + 1):
        survivor = survivor.merge(row.cells[index])
    return survivor


def table_borders_none(table) -> None:
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        borders.append(el(f"w:{edge}", val="nil"))
    replace_child(table._tbl.tblPr, "w:tblBorders", borders)


def fixed_layout(table) -> None:
    """Pin column widths. Without this Word re-flows columns to fit content."""
    replace_child(table._tbl.tblPr, "w:tblLayout", el("w:tblLayout", type="fixed"))
    table.autofit = False


def table_caption(table, caption: str, description: str = "") -> None:
    """Tag a table as a data table.

    ``w:tblCaption`` and ``w:tblDescription`` are announced by assistive
    technology, and they give tooling a reliable way to distinguish a data table
    (which must repeat its header row) from a layout grid (which has no header).
    """
    replace_child(table._tbl.tblPr, "w:tblCaption", el("w:tblCaption", val=caption))
    if description:
        replace_child(table._tbl.tblPr, "w:tblDescription",
                      el("w:tblDescription", val=description))


def table_indent(table, dxa: int) -> None:
    replace_child(table._tbl.tblPr, "w:tblInd", el("w:tblInd", w=dxa, type="dxa"))


def table_cell_spacing(table, dxa: int) -> None:
    replace_child(table._tbl.tblPr, "w:tblCellSpacing",
                  el("w:tblCellSpacing", w=dxa, type="dxa"))


def row_height(row, mm: float, exact: bool = False) -> None:
    """Minimum (or exact) row height. Always prefer ``atLeast`` for flow safety:
    an ``exact`` row silently clips content that grows past it."""
    replace_child(row._tr.get_or_add_trPr(), "w:trHeight",
                  el("w:trHeight", val=int(Mm(mm).twips),
                     hRule="exact" if exact else "atLeast"))


def keep_row_together(row) -> None:
    """Stop a row splitting across a page break."""
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        splice(tr_pr, OxmlElement("w:cantSplit"))


def repeat_header(row) -> None:
    """Repeat this row at the top of every page the table spans."""
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        splice(tr_pr, OxmlElement("w:tblHeader"))


# ==========================================================================
# Paragraphs and runs
# ==========================================================================

def para_format(
    paragraph,
    *,
    before: float = 0,
    after: float = 0,
    line: float | None = None,
    align=None,
    keep_with_next: bool = False,
    keep_lines: bool = True,
    left_indent: float | None = None,
    right_indent: float | None = None,
    hanging: float | None = None,
    first_line: float | None = None,
    page_break_before: bool = False,
) -> None:
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line is not None:
        pf.line_spacing = line
    if align is not None:
        paragraph.alignment = align
    pf.keep_together = keep_lines
    pf.keep_with_next = keep_with_next
    pf.page_break_before = page_break_before
    if left_indent is not None:
        pf.left_indent = Pt(left_indent)
    if right_indent is not None:
        pf.right_indent = Pt(right_indent)
    if hanging is not None:
        pf.first_line_indent = Pt(-hanging)
    elif first_line is not None:
        pf.first_line_indent = Pt(first_line)
    # Widow/orphan control: never leave one line of a paragraph stranded.
    p_pr = paragraph._p.get_or_add_pPr()
    replace_child(p_pr, "w:widowControl", el("w:widowControl", val="1"))


def write(
    paragraph,
    text: str,
    *,
    font: str = "Calibri",
    size: float = 10.0,
    bold: bool = False,
    italic: bool = False,
    colour: str = "000000",
    caps: bool = False,
    small_caps: bool = False,
    tracking: float = 0.0,
    underline: bool = False,
    superscript: bool = False,
    highlight: str | None = None,
):
    """Append a styled run. ``tracking`` is extra character spacing in points."""
    run = paragraph.add_run(text)
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.underline = underline
    run.font.color.rgb = RGBColor.from_string(colour)
    r_pr = run._r.get_or_add_rPr()
    # Set every script slot so the face survives a locale switch.
    fonts = get_or_add(r_pr, "w:rFonts")
    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
        fonts.set(qn(f"w:{attr}"), font)
    if caps:
        replace_child(r_pr, "w:caps", el("w:caps", val="1"))
    if small_caps:
        replace_child(r_pr, "w:smallCaps", el("w:smallCaps", val="1"))
    if tracking:
        replace_child(r_pr, "w:spacing", el("w:spacing", val=int(round(tracking * 20))))
    if superscript:
        replace_child(r_pr, "w:vertAlign", el("w:vertAlign", val="superscript"))
    if highlight:
        replace_child(r_pr, "w:shd", el("w:shd", val="clear", color="auto", fill=highlight))
    return run


_PARA_KEYS = {
    "before", "after", "line", "align", "keep_with_next", "keep_lines",
    "left_indent", "right_indent", "hanging", "first_line", "page_break_before",
}


def para(container, text: str = "", **kwargs):
    """Add a paragraph to a document, cell, header or footer, and optionally
    write one run into it. Paragraph-level and run-level keywords are separated
    automatically, so callers write a single flat call."""
    p_kwargs = {k: v for k, v in kwargs.items() if k in _PARA_KEYS}
    r_kwargs = {k: v for k, v in kwargs.items() if k not in _PARA_KEYS}
    paragraph = container.add_paragraph()
    para_format(paragraph, **p_kwargs)
    if text:
        write(paragraph, text, **r_kwargs)
    return paragraph


def tab_stop(paragraph, position_pt: float, alignment: str = "left",
             leader: str | None = None) -> None:
    tabs = get_or_add(paragraph._p.get_or_add_pPr(), "w:tabs")
    node = el("w:tab", val=alignment, pos=int(position_pt * 20))
    if leader:
        node.set(qn("w:leader"), leader)
    tabs.append(node)


def clear(cell):
    """Remove the empty paragraph python-docx seeds into every new cell."""
    for child in list(cell._tc):
        if child.tag == qn("w:p"):
            cell._tc.remove(child)
    return cell


def horizontal_rule(paragraph, colour: str, size: int = 8, space: int = 1) -> None:
    """Draw a bottom border on a paragraph — the cheapest reliable divider."""
    p_bdr = get_or_add(paragraph._p.get_or_add_pPr(), "w:pBdr")
    replace_child(p_bdr, "w:bottom",
                  el("w:bottom", val="single", sz=size, space=space, color=colour))


# ==========================================================================
# Fields, breaks and links
# ==========================================================================

def page_field(paragraph, instruction: str, **run_kwargs) -> None:
    """Insert a Word field such as ``PAGE``, ``NUMPAGES`` or ``TOC \\o "1-3"``."""
    run = write(paragraph, "", **run_kwargs)
    begin = el("w:fldChar", fldCharType="begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {instruction} "
    separate = el("w:fldChar", fldCharType="separate")
    end = el("w:fldChar", fldCharType="end")
    for node in (begin, instr, separate, end):
        run._r.append(node)


def bookmark(paragraph, name: str, bookmark_id: int) -> None:
    """Named anchor, so a generated table of contents can link to it."""
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def hyperlink(paragraph, url: str, text: str, colour: str = "1D6FE0",
              **run_kwargs) -> None:
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), r_id)
    holder = paragraph.add_run()
    run = write(paragraph, text, colour=colour, underline=True, **run_kwargs)
    link.append(run._r)
    holder._r.addnext(link)
    holder._r.getparent().remove(holder._r)


def page_break(container) -> None:
    paragraph = container.add_paragraph()
    para_format(paragraph, before=0, after=0)
    paragraph.add_run().add_break(WD_BREAK.PAGE)


def spacer(container, points: float = 6) -> None:
    """Vertical whitespace that cannot collapse. A zero-height paragraph with a
    fixed line height is the only spacing Word, LibreOffice and Google Docs all
    agree on."""
    paragraph = container.add_paragraph()
    para_format(paragraph, before=0, after=0)
    paragraph.paragraph_format.line_spacing = Pt(points)
    paragraph.add_run().font.size = Pt(1)


# ==========================================================================
# Table scaffolding
# ==========================================================================

def new_table(container, rows: int, cols: int, widths_mm: list[float],
              align=WD_TABLE_ALIGNMENT.CENTER):
    """Fixed-layout, borderless table sized in millimetres.

    ``container`` may be the document, a cell, or a header/footer — the latter
    two require an explicit total width, which ``Document.add_table`` rejects.
    """
    total = Mm(sum(widths_mm))
    try:
        table = container.add_table(rows=rows, cols=cols, width=total)
    except TypeError:
        table = container.add_table(rows=rows, cols=cols)
    table.alignment = align
    table_borders_none(table)
    fixed_layout(table)
    table_indent(table, 0)
    for row in table.rows:
        keep_row_together(row)
        for index, cell in enumerate(row.cells):
            cell.width = Mm(widths_mm[index])
    for index, column in enumerate(table.columns):
        column.width = Mm(widths_mm[index])
    return table


def single_cell(container, width_mm: float, *, fill: str | None = None,
                border=None, pad: tuple[int, int] = (170, 200),
                accent_bar=None, accent_edge: str = "left"):
    """A one-cell panel — the workhorse for cards, callouts and banners.

    Returns the cell, cleared and ready for paragraphs. ``accent_bar`` draws a
    heavy rule on ``accent_edge`` and overrides ``border`` on that edge.
    """
    table = new_table(container, 1, 1, [width_mm])
    cell = clear(table.cell(0, 0))
    if fill:
        shade(cell, fill)
    cell_margins(cell, pad[0], pad[1], pad[0], pad[1])
    edges = {"top": border, "left": border, "bottom": border, "right": border}
    if accent_bar:
        edges[accent_edge] = accent_bar
    cell_borders(cell, **edges)
    return cell


def set_style_font(style, font: str, size: float, colour: str) -> None:
    """Point a paragraph style at a face for every script slot."""
    style.font.name = font
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(colour)
    fonts = get_or_add(style.element.get_or_add_rPr(), "w:rFonts")
    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
        fonts.set(qn(f"w:{attr}"), font)
